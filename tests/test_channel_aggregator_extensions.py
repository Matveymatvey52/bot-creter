"""Extensions for the channel_aggregator template — schema presets, 50-channel
limit, extract/views/forwards columns, report_schedules, report generation
(CSV/docx/Sheets), and the standalone templates/channel_aggregator.py wrapper.

See docs/CHANNEL_AGGREGATOR_TEMPLATE_DESIGN.md §5 for the owner decisions
this implements. DB-level tests reuse tests/conftest.py's isolated_db
fixture, same convention as tests/test_channel_monitor_db.py. Gemini/Sheets/
Telethon are mocked — no real network calls.
"""
from __future__ import annotations

import json
from unittest.mock import AsyncMock, patch

import pytest

from db.database import (
    add_channel_post_with_metrics,
    add_monitored_channel,
    count_monitored_channels,
    deactivate_report_schedule,
    get_active_report_schedule,
    get_all_active_report_schedules,
    get_monitored_channel,
    get_posts_for_report,
    get_posts_missing_extraction,
    init_channel_monitor_tables,
    set_channel_post_extracted,
    set_monitored_channel_schema,
    set_report_schedule,
    set_report_schedule_last_sent,
)
from templates import channel_aggregator

BOT_A = 111
BOT_B = 222


# ── DB layer ──────────────────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_count_monitored_channels(isolated_db):
    await init_channel_monitor_tables(isolated_db)
    assert await count_monitored_channels(isolated_db, BOT_A) == 0
    await add_monitored_channel(isolated_db, BOT_A, "a", -1, "A")
    await add_monitored_channel(isolated_db, BOT_A, "b", -2, "B")
    await add_monitored_channel(isolated_db, BOT_B, "c", -3, "C")
    assert await count_monitored_channels(isolated_db, BOT_A) == 2
    assert await count_monitored_channels(isolated_db, BOT_B) == 1


@pytest.mark.asyncio
async def test_set_monitored_channel_schema(isolated_db):
    await init_channel_monitor_tables(isolated_db)
    channel_id = await add_monitored_channel(isolated_db, BOT_A, "chan", -1, "Chan")
    row = await get_monitored_channel(isolated_db, channel_id)
    assert row["extract_schema"] is None

    fields = json.dumps(["должность", "зарплата"], ensure_ascii=False)
    await set_monitored_channel_schema(isolated_db, channel_id, fields)
    row = await get_monitored_channel(isolated_db, channel_id)
    assert json.loads(row["extract_schema"]) == ["должность", "зарплата"]


@pytest.mark.asyncio
async def test_add_channel_post_with_metrics_stores_views_forwards(isolated_db):
    await init_channel_monitor_tables(isolated_db)
    channel_id = await add_monitored_channel(isolated_db, BOT_A, "chan", -1, "Chan")
    post_id = await add_channel_post_with_metrics(isolated_db, channel_id, 1, "hello", views=100, forwards=5)
    posts = await get_posts_for_report(isolated_db, BOT_A)
    post = next(p for p in posts if p["id"] == post_id)
    assert post["views"] == 100
    assert post["forwards"] == 5


@pytest.mark.asyncio
async def test_get_posts_missing_extraction_only_schema_channels(isolated_db):
    await init_channel_monitor_tables(isolated_db)
    schema_channel = await add_monitored_channel(isolated_db, BOT_A, "s", -1, "S")
    await set_monitored_channel_schema(isolated_db, schema_channel, json.dumps(["поле"]))
    plain_channel = await add_monitored_channel(isolated_db, BOT_A, "p", -2, "P")

    p1 = await add_channel_post_with_metrics(isolated_db, schema_channel, 1, "with schema")
    await add_channel_post_with_metrics(isolated_db, plain_channel, 1, "no schema")

    missing = await get_posts_missing_extraction(isolated_db)
    ids = {p["id"] for p in missing}
    assert p1 in ids
    assert len(missing) == 1

    await set_channel_post_extracted(isolated_db, p1, json.dumps({"поле": "значение"}, ensure_ascii=False))
    missing_after = await get_posts_missing_extraction(isolated_db)
    assert not any(p["id"] == p1 for p in missing_after)


@pytest.mark.asyncio
async def test_get_posts_for_report_filters_by_channel_and_since(isolated_db):
    await init_channel_monitor_tables(isolated_db)
    chan_a = await add_monitored_channel(isolated_db, BOT_A, "a", -1, "A")
    chan_b = await add_monitored_channel(isolated_db, BOT_A, "b", -2, "B")
    await add_channel_post_with_metrics(isolated_db, chan_a, 1, "post a")
    await add_channel_post_with_metrics(isolated_db, chan_b, 1, "post b")

    only_a = await get_posts_for_report(isolated_db, BOT_A, channel_row_id=chan_a)
    assert {p["text"] for p in only_a} == {"post a"}

    future_since = "2999-01-01 00:00:00"
    none_found = await get_posts_for_report(isolated_db, BOT_A, since=future_since)
    assert none_found == []


@pytest.mark.asyncio
async def test_report_schedule_crud(isolated_db):
    await init_channel_monitor_tables(isolated_db)
    assert await get_active_report_schedule(isolated_db, BOT_A) is None

    sched_id = await set_report_schedule(isolated_db, BOT_A, "daily", hour=9, minute=0)
    active = await get_active_report_schedule(isolated_db, BOT_A)
    assert active["id"] == sched_id
    assert active["frequency"] == "daily"

    # setting a new schedule deactivates the old one (one active per bot)
    new_id = await set_report_schedule(isolated_db, BOT_A, "weekly", hour=10, minute=30, weekday=0)
    active = await get_active_report_schedule(isolated_db, BOT_A)
    assert active["id"] == new_id
    assert active["weekday"] == 0

    all_active = await get_all_active_report_schedules(isolated_db)
    assert len(all_active) == 1

    await set_report_schedule_last_sent(isolated_db, new_id, "2026-08-19 10:30:00")
    active = await get_active_report_schedule(isolated_db, BOT_A)
    assert active["last_sent_at"] == "2026-08-19 10:30:00"

    await deactivate_report_schedule(isolated_db, BOT_A)
    assert await get_active_report_schedule(isolated_db, BOT_A) is None


# ── report generation (features/channel_monitor.py) ─────────────────────────

@pytest.mark.asyncio
async def test_generate_report_rows_uses_schema_columns(isolated_db):
    from features.channel_monitor import generate_report_rows

    await init_channel_monitor_tables(isolated_db)
    channel_id = await add_monitored_channel(isolated_db, BOT_A, "jobs", -1, "Jobs")
    await set_monitored_channel_schema(isolated_db, channel_id, json.dumps(["должность", "зарплата"], ensure_ascii=False))
    post_id = await add_channel_post_with_metrics(isolated_db, channel_id, 1, "raw text", views=10, forwards=1)
    await set_channel_post_extracted(
        isolated_db, post_id, json.dumps({"должность": "Разработчик", "зарплата": "200000"}, ensure_ascii=False)
    )

    headers, rows = await generate_report_rows(isolated_db, BOT_A)
    assert headers == ["Канал", "Дата", "должность", "зарплата", "Просмотры", "Репосты"]
    assert rows[0][0] == "Jobs"
    assert "Разработчик" in rows[0]
    assert "200000" in rows[0]


@pytest.mark.asyncio
async def test_generate_report_rows_generic_without_schema(isolated_db):
    from features.channel_monitor import generate_report_rows

    await init_channel_monitor_tables(isolated_db)
    channel_id = await add_monitored_channel(isolated_db, BOT_A, "news", -1, "News")
    await add_channel_post_with_metrics(isolated_db, channel_id, 1, "plain text post")

    headers, rows = await generate_report_rows(isolated_db, BOT_A)
    assert headers == ["Канал", "Дата", "Summary/текст", "Просмотры", "Репосты"]
    assert "plain text post" in rows[0]


@pytest.mark.asyncio
async def test_generate_report_rows_empty(isolated_db):
    from features.channel_monitor import generate_report_rows

    await init_channel_monitor_tables(isolated_db)
    headers, rows = await generate_report_rows(isolated_db, BOT_A)
    assert headers == []
    assert rows == []


def test_build_csv_bytes():
    from features.channel_monitor import build_csv_bytes

    content = build_csv_bytes(["A", "B"], [["1", "2"], ["3", "4"]])
    text = content.decode("utf-8-sig")
    assert "A,B" in text
    assert "1,2" in text


def test_build_docx_bytes_produces_valid_document():
    from docx import Document

    from features.channel_monitor import build_docx_bytes
    import io

    content = build_docx_bytes(["A", "B"], [["1", "2"]], title="Test Report")
    doc = Document(io.BytesIO(content))
    assert doc.tables
    table = doc.tables[0]
    assert table.rows[0].cells[0].text == "A"
    assert table.rows[1].cells[1].text == "2"


@pytest.mark.asyncio
async def test_sync_report_to_sheets_calls_write_row():
    from features.channel_monitor import sync_report_to_sheets

    with patch("features.sheets.write_row", new=AsyncMock()) as mock_write:
        n = await sync_report_to_sheets(bot_id=123, headers=["A", "B"], rows=[["1", "2"], ["3", "4"]])
    assert n == 2
    assert mock_write.await_count == 3  # header + 2 rows


@pytest.mark.asyncio
async def test_sync_report_to_sheets_propagates_no_sheet_connected():
    from features.channel_monitor import sync_report_to_sheets

    with patch("features.sheets.write_row", new=AsyncMock(side_effect=ValueError("no sheet"))):
        with pytest.raises(ValueError):
            await sync_report_to_sheets(bot_id=123, headers=["A"], rows=[["1"]])


# ── services/gemini_service.extract_structured ───────────────────────────────

@pytest.mark.asyncio
async def test_extract_structured_returns_none_without_api_key(monkeypatch):
    from services import gemini_service

    monkeypatch.setattr(gemini_service, "GEMINI_API_KEY", "")
    result = await gemini_service.extract_structured("some post text", ["field1"])
    assert result is None


@pytest.mark.asyncio
async def test_extract_structured_returns_none_for_empty_text():
    from services import gemini_service

    assert await gemini_service.extract_structured("", ["field1"]) is None
    assert await gemini_service.extract_structured("text", []) is None


@pytest.mark.asyncio
async def test_extract_structured_parses_valid_json(monkeypatch):
    from services import gemini_service

    monkeypatch.setattr(gemini_service, "GEMINI_API_KEY", "fake-key")

    class FakeResponse:
        text = '{"должность": "Аналитик", "зарплата": "150000"}'

    fake_models = AsyncMock()
    fake_models.generate_content = AsyncMock(return_value=FakeResponse())
    fake_client = type("C", (), {"aio": type("Aio", (), {"models": fake_models})()})()
    monkeypatch.setattr(gemini_service, "_get_client", lambda: fake_client)

    result = await gemini_service.extract_structured("some post", ["должность", "зарплата"])
    assert result == {"должность": "Аналитик", "зарплата": "150000"}


@pytest.mark.asyncio
async def test_extract_structured_strips_markdown_fence(monkeypatch):
    from services import gemini_service

    monkeypatch.setattr(gemini_service, "GEMINI_API_KEY", "fake-key")

    class FakeResponse:
        text = '```json\n{"a": "b"}\n```'

    fake_models = AsyncMock()
    fake_models.generate_content = AsyncMock(return_value=FakeResponse())
    fake_client = type("C", (), {"aio": type("Aio", (), {"models": fake_models})()})()
    monkeypatch.setattr(gemini_service, "_get_client", lambda: fake_client)

    result = await gemini_service.extract_structured("post", ["a"])
    assert result == {"a": "b"}


@pytest.mark.asyncio
async def test_extract_structured_returns_none_on_malformed_json(monkeypatch):
    from services import gemini_service

    monkeypatch.setattr(gemini_service, "GEMINI_API_KEY", "fake-key")

    class FakeResponse:
        text = "not json at all"

    fake_models = AsyncMock()
    fake_models.generate_content = AsyncMock(return_value=FakeResponse())
    fake_client = type("C", (), {"aio": type("Aio", (), {"models": fake_models})()})()
    monkeypatch.setattr(gemini_service, "_get_client", lambda: fake_client)

    result = await gemini_service.extract_structured("post", ["a"])
    assert result is None


# ── mini-app config ──────────────────────────────────────────────────────
# miniapp_config's declared table/field names must match init_channel_monitor
# _tables' real schema — miniapp_api.py builds SQL directly off these names,
# so a drift here would 500 at request time instead of failing a test.

def test_miniapp_config_resource_names():
    names = {r["name"] for r in channel_aggregator.miniapp_config["resources"]}
    assert names == {"channels", "posts"}


def test_channels_resource_targets_monitored_channels_table():
    channels = next(r for r in channel_aggregator.miniapp_config["resources"] if r["name"] == "channels")
    assert channels["table"] == "monitored_channels"
    assert channels["creatable"] is False
    assert "channel_title" in {f["name"] for f in channels["fields"]}


def test_posts_resource_targets_channel_posts_table():
    posts = next(r for r in channel_aggregator.miniapp_config["resources"] if r["name"] == "posts")
    assert posts["table"] == "channel_posts"
    assert posts["creatable"] is False
    field_names = {f["name"] for f in posts["fields"]}
    assert field_names == {"monitored_channel_id", "text", "summary", "views", "forwards", "posted_at"}


@pytest.mark.asyncio
async def test_miniapp_config_fields_match_real_schema(isolated_db):
    await channel_aggregator.init_db(isolated_db)
    import aiosqlite
    async with aiosqlite.connect(isolated_db) as db:
        for resource in channel_aggregator.miniapp_config["resources"]:
            cur = await db.execute(f"PRAGMA table_info({resource['table']})")
            real_columns = {row[1] for row in await cur.fetchall()}
            declared = {f["name"] for f in resource["fields"]} | {"id"}
            assert declared.issubset(real_columns), (
                f"{resource['name']}: declared fields {declared} not all in "
                f"real columns {real_columns}"
            )
