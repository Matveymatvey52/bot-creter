"""features/word_export.py tests — the word-export feature ТЗ's own test list:
document creation with text, table/field insertion, and export without errors.

Run with: python -m unittest tests.test_word_export_module
"""
from __future__ import annotations

import io
import unittest

from docx import Document

import features.word_export as word_export_module


class BuildDocumentTests(unittest.TestCase):
    def test_returns_nonempty_docx_bytes(self):
        data = word_export_module.build_document("Анкета", [("Имя", "Иван")])
        self.assertIsInstance(data, bytes)
        self.assertGreater(len(data), 0)

    def test_document_is_valid_and_openable(self):
        data = word_export_module.build_document("Отчёт", [("Поле", "Значение")])
        document = Document(io.BytesIO(data))
        self.assertEqual(document.paragraphs[0].text, "Отчёт")

    def test_fields_render_as_table_rows(self):
        fields = [("ФИО", "Иван Иванов"), ("Email", "test@example.com")]
        data = word_export_module.build_document("Анкета", fields)
        document = Document(io.BytesIO(data))
        table = document.tables[0]
        self.assertEqual(len(table.rows), 2)
        self.assertEqual(table.rows[0].cells[0].text, "ФИО")
        self.assertEqual(table.rows[0].cells[1].text, "Иван Иванов")

    def test_empty_fields_list_does_not_raise(self):
        data = word_export_module.build_document("Пустой отчёт", [])
        document = Document(io.BytesIO(data))
        self.assertEqual(len(document.tables), 0)

    def test_metadata_footer_included_by_default(self):
        data = word_export_module.build_document(
            "Анкета", [("Имя", "Иван")], bot_name="test_bot", user_name="Иван"
        )
        document = Document(io.BytesIO(data))
        full_text = "\n".join(p.text for p in document.paragraphs)
        self.assertIn("test_bot", full_text)
        self.assertIn("Иван", full_text)

    def test_metadata_footer_omitted_when_disabled(self):
        data = word_export_module.build_document(
            "Анкета", [("Имя", "Иван")], include_metadata=False, bot_name="test_bot"
        )
        document = Document(io.BytesIO(data))
        full_text = "\n".join(p.text for p in document.paragraphs)
        self.assertNotIn("test_bot", full_text)

    def test_router_exported_for_registry_convention(self):
        self.assertTrue(hasattr(word_export_module, "router"))

    def test_non_string_field_values_are_coerced(self):
        data = word_export_module.build_document("Отчёт", [("Возраст", 30), ("Активен", True)])
        document = Document(io.BytesIO(data))
        table = document.tables[0]
        self.assertEqual(table.rows[0].cells[1].text, "30")
        self.assertEqual(table.rows[1].cells[1].text, "True")

    def test_long_field_value_does_not_raise(self):
        long_value = "А" * 5000
        data = word_export_module.build_document("Отчёт", [("Комментарий", long_value)])
        document = Document(io.BytesIO(data))
        self.assertEqual(document.tables[0].rows[0].cells[1].text, long_value)


if __name__ == "__main__":
    unittest.main()
