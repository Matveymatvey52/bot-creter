"""services/attachment_service.py — text extraction from PDF/Word/Excel
attachments uploaded during /create's requirements gathering, plus the
image/text content-block builders used to assemble Claude vision messages.

Run with: python -m unittest tests.test_attachment_service
"""
from __future__ import annotations

import io
import unittest

from services import attachment_service as svc


class ImageAndTextBlockTests(unittest.TestCase):
    def test_build_image_block_base64_encodes(self):
        block = svc.build_image_block(b"\xff\xd8raw", "image/jpeg")
        self.assertEqual(block["type"], "image")
        self.assertEqual(block["source"]["media_type"], "image/jpeg")
        self.assertEqual(block["source"]["type"], "base64")
        self.assertIsInstance(block["source"]["data"], str)

    def test_build_text_block(self):
        self.assertEqual(svc.build_text_block("hi"), {"type": "text", "text": "hi"})

    def test_truncate_long_text(self):
        long_text = "a" * (svc.MAX_EXTRACTED_CHARS + 500)
        truncated = svc._truncate(long_text)
        self.assertLessEqual(len(truncated), svc.MAX_EXTRACTED_CHARS + len("\n…[обрезано]"))
        self.assertTrue(truncated.endswith("[обрезано]"))


class PdfExtractionTests(unittest.TestCase):
    def _make_pdf(self) -> bytes:
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        buf = io.BytesIO()
        writer.write(buf)
        return buf.getvalue()

    def test_extract_pdf_text_runs_without_raising(self):
        # Building a PDF with real extractable text needs reportlab, which
        # isn't a project dependency. This verifies extract_pdf_text runs
        # end-to-end on a minimal valid (textless) PDF and returns a string
        # (possibly empty) rather than crashing.
        result = svc.extract_pdf_text(self._make_pdf())
        self.assertIsInstance(result, str)

    def test_extract_document_text_routes_pdf_extension(self):
        result = svc.extract_document_text(self._make_pdf(), "brief.PDF")
        self.assertIsInstance(result, str)


class DocxExtractionTests(unittest.TestCase):
    def _make_docx(self) -> bytes:
        from docx import Document

        doc = Document()
        doc.add_paragraph("Hello from docx")
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Name"
        table.rows[0].cells[1].text = "Alice"
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_extract_docx_text_includes_paragraphs_and_tables(self):
        text = svc.extract_docx_text(self._make_docx())
        self.assertIn("Hello from docx", text)
        self.assertIn("Name", text)
        self.assertIn("Alice", text)

    def test_extract_document_text_routes_docx_extension(self):
        text = svc.extract_document_text(self._make_docx(), "notes.docx")
        self.assertIn("Hello from docx", text)


class XlsxExtractionTests(unittest.TestCase):
    def _make_xlsx(self) -> bytes:
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "Orders"
        ws.append(["Order", "Amount"])
        ws.append(["A1", 100])
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    def test_extract_xlsx_text_includes_sheet_name_and_rows(self):
        text = svc.extract_xlsx_text(self._make_xlsx())
        self.assertIn("Orders", text)
        self.assertIn("Order", text)
        self.assertIn("A1", text)
        self.assertIn("100", text)

    def test_extract_document_text_routes_xlsx_extension(self):
        text = svc.extract_document_text(self._make_xlsx(), "orders.xlsx")
        self.assertIn("Orders", text)


class UnsupportedExtensionTests(unittest.TestCase):
    def test_unknown_extension_returns_none(self):
        self.assertIsNone(svc.extract_document_text(b"whatever", "archive.zip"))

    def test_no_filename_returns_none(self):
        self.assertIsNone(svc.extract_document_text(b"whatever", ""))


if __name__ == "__main__":
    unittest.main()
