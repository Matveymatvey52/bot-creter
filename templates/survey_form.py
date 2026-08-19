# TEMPLATE: survey_form
# USE FOR: анкета/опросник для клиентов — админ загружает анкету (текст, голос, PDF/Word/Excel-файл), бот парсит вопросы и подтверждает их; клиенты проходят анкету по вопросу за раз, ответы сохраняются в таблицу и экспортируются в Word/Excel по запросу; администратор получает уведомление о каждой заполненной анкете
# CUSTOMIZE: sections marked with # CUSTOMIZE
from __future__ import annotations

import asyncio
import html
import io
import json
import logging
import os
import re
import time
from dataclasses import dataclass
from pathlib import Path

import aiohttp
import aiosqlite
from aiogram import BaseMiddleware, Bot, Dispatcher, F, Router
from aiogram.exceptions import TelegramBadRequest
from aiogram.filters import Command
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.types import (
    BufferedInputFile, CallbackQuery, Document, FSInputFile, InlineKeyboardButton,
    InlineKeyboardMarkup, Message,
)
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter
from pypdf import PdfReader

from config import ANTHROPIC_API_KEY, ASSEMBLYAI_API_KEY
from features.word_export import build_document

# ── CUSTOMIZE ────────────────────────────────────────────────────────────────
# Same status as every other template's CUSTOMIZE block: per-file source-text
# customization Claude edits when generating a specific bot, not per-bot
# runtime state (that's config.db_path/admins_file below). Surveys and
# responses themselves are runtime data, created via the admin upload flow
# and the client answer flow.
BOT_DESCRIPTION = "Анкета для клиентов: загрузите вопросы (текст/голос/файл), клиенты проходят опрос, ответы — в таблице и экспорте."
WELCOME_TEXT = (
    "📋 <b>Анкеты</b>\n\n"
    "Загрузите анкету, следите за ответами клиентов и выгружайте их в Word/Excel.\n\n"
    "Выберите действие:"
)
CLIENT_WELCOME_TEXT = (
    "📋 Здравствуйте!\n\n"
    "Ответьте, пожалуйста, на несколько вопросов:"
)
UPLOAD_PROMPT_TEXT = (
    "📥 <b>Загрузка анкеты</b>\n\n"
    "Пришлите вопросы одним из способов:\n"
    "• текстом — каждый вопрос с новой строки или пронумерованный список\n"
    "• голосовым сообщением\n"
    "• файлом — PDF, Word (.docx) или Excel (.xlsx)"
)
THANK_YOU_TEXT = "🙏 Спасибо! Ваши ответы сохранены."
MAX_QUESTIONS = 30                # CUSTOMIZE: верхняя граница длины анкеты
QUESTION_MAX_LEN = 500
ANSWER_MAX_LEN = 1000
FILE_MAX_SIZE_BYTES = 15 * 1024 * 1024   # CUSTOMIZE: верхняя граница загружаемого файла анкеты
VOICE_LANGUAGE_CODE = "ru"        # CUSTOMIZE
RECENT_RESPONSES_LIMIT = 10
# ── END CUSTOMIZE ─────────────────────────────────────────────────────────────

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
router = Router()

# Same staleness guard as templates/feedback_survey.py's FLOW_TIMEOUT_SECONDS/
# _flow_expired: without it a client who starts a questionnaire and goes quiet
# for a long time has their next unrelated message silently swallowed into a
# half-finished flow from a stale session.
FLOW_TIMEOUT_SECONDS = 600


def _flow_expired(data: dict) -> bool:
    started_at = data.get("started_at")
    return started_at is None or (time.time() - started_at) > FLOW_TIMEOUT_SECONDS


async def _edit_text_safe(message: Message, text: str, **kwargs) -> None:
    """Two near-simultaneous taps of the same button can both pass the FSM
    state filter before either has advanced the state (MemoryStorage reads
    don't suspend), so the loser redundantly repeats the identical
    edit_text call — Telegram rejects that with "message is not modified".
    The state itself was already correctly advanced by the winner, so this
    is harmless; swallow only that specific error, anything else re-raises."""
    try:
        await message.edit_text(text, **kwargs)
    except TelegramBadRequest as e:
        if "message is not modified" not in str(e):
            raise


# ── config ───────────────────────────────────────────────────────────────────
# Same pattern as every other template — see docs/STAGE2_DESIGN.md.

@dataclass
class SurveyFormConfig:
    bot_name: str
    db_path: str
    admins_file: Path
    welcome_image: Path
    display_name: str | None = None
    group_chat_id: str | None = None
    bot_id: int | None = None


def _paths_for(name: str, data_dir: Path) -> SurveyFormConfig:
    return SurveyFormConfig(
        bot_name=name,
        db_path=str(data_dir / f"{name}_data.db"),
        admins_file=data_dir / f"admins_{name}.json",
        welcome_image=data_dir / "bot_images" / f"{name}.jpg",
    )


def config_from_env() -> SurveyFormConfig:
    """Standalone/subprocess mode."""
    name = Path(__file__).stem
    data_dir = Path(os.getenv("DATA_DIR", "./data"))
    data_dir.mkdir(exist_ok=True)
    return _paths_for(name, data_dir)


def config_from_bot_row(bot_row: dict, data_dir: Path) -> SurveyFormConfig:
    """Webhook runtime mode. Paths built from bot_row["bot_id"] (bots.id, the
    physically unique AUTOINCREMENT PK) — NOT bot_row["name"] — same reasoning
    as every other template's config_from_bot_row (see docs/STAGE2_DESIGN.md
    "Изоляция по bots.id")."""
    bot_id = bot_row["bot_id"]
    config = SurveyFormConfig(
        bot_name=bot_row["name"],
        db_path=str(data_dir / f"bot_{bot_id}_data.db"),
        admins_file=data_dir / f"admins_{bot_id}.json",
        welcome_image=data_dir / "bot_images" / f"bot_{bot_id}.jpg",
    )
    config.display_name = bot_row.get("display_name")
    config.group_chat_id = bot_row.get("group_chat_id")
    config.bot_id = bot_id
    return config


class ConfigMiddleware(BaseMiddleware):
    """Injects this bot's SurveyFormConfig into data["config"]."""

    def __init__(self, config: SurveyFormConfig) -> None:
        self.config = config
        super().__init__()

    async def __call__(self, handler, event, data):
        data["config"] = self.config
        return await handler(event, data)


# ── admin helpers ─────────────────────────────────────────────────────────────

def _load_admins(admins_file: Path) -> set:
    try:
        return set(json.loads(admins_file.read_text()).get("ids", []))
    except Exception:
        return set()

def _save_admins(admins_file: Path, ids: set) -> None:
    admins_file.write_text(json.dumps({"ids": list(ids)}, ensure_ascii=False))

def _is_admin(user_id: int, config: SurveyFormConfig) -> bool:
    return str(user_id) in _load_admins(config.admins_file)


def _valid_admin_id(text: str) -> bool:
    """Same guard as templates/moderator.py's _valid_admin_id() — rejects
    unbounded-length strings, non-ASCII look-alike digits, and 0/leading-zero
    phantom ids that would inflate the admin list without ever matching
    _is_admin()."""
    if not (bool(text) and text.isascii() and text.isdigit() and len(text) <= 15):
        return False
    return int(text) > 0 and str(int(text)) == text


def _esc(value, max_len: int = 500) -> str:
    """HTML-escapes AND length-bounds any user-supplied text before it goes into
    a parse_mode="HTML" message — same helper/rationale as every other
    template's _esc()."""
    text = str(value) if value is not None else ""
    if len(text) > max_len:
        text = text[:max_len] + "…"
    return html.escape(text)


def _join_bounded(lines: list[str], limit: int = 3500) -> str:
    """Joins lines with a length budget, dropping only WHOLE trailing lines —
    same rationale as every other template's _join_bounded()."""
    out: list[str] = []
    total = 0
    for line in lines:
        if total + len(line) + 1 > limit:
            out.append("…")
            break
        out.append(line)
        total += len(line) + 1
    return "\n".join(out)


# ── db ────────────────────────────────────────────────────────────────────────

async def init_db(db_path: str):
    async with aiosqlite.connect(db_path) as db:
        await db.execute("""
            CREATE TABLE IF NOT EXISTS surveys (
                id         INTEGER PRIMARY KEY AUTOINCREMENT,
                title      TEXT NOT NULL,
                is_active  INTEGER NOT NULL DEFAULT 1,
                created_at TEXT NOT NULL DEFAULT (datetime('now','localtime'))
            )
        """)
        await db.execute("""
            CREATE TABLE IF NOT EXISTS survey_questions (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                survey_id   INTEGER NOT NULL REFERENCES surveys(id),
                position    INTEGER NOT NULL,
                question    TEXT NOT NULL
            )
        """)
        await db.execute("""
            CREATE TABLE IF NOT EXISTS survey_responses (
                id               INTEGER PRIMARY KEY AUTOINCREMENT,
                survey_id        INTEGER NOT NULL REFERENCES surveys(id),
                user_telegram_id INTEGER NOT NULL,
                user_name        TEXT,
                submitted_at     TEXT NOT NULL DEFAULT (datetime('now','localtime'))
            )
        """)
        await db.execute("""
            CREATE TABLE IF NOT EXISTS survey_answers (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                response_id INTEGER NOT NULL REFERENCES survey_responses(id),
                question_id INTEGER NOT NULL REFERENCES survey_questions(id),
                answer      TEXT NOT NULL
            )
        """)
        await db.execute("CREATE INDEX IF NOT EXISTS idx_survey_questions_survey ON survey_questions(survey_id, position)")
        await db.execute("CREATE INDEX IF NOT EXISTS idx_survey_responses_survey ON survey_responses(survey_id)")
        await db.execute("CREATE INDEX IF NOT EXISTS idx_survey_answers_response ON survey_answers(response_id)")
        await db.commit()


async def _create_survey(db_path: str, title: str, questions: list[str]) -> int:
    async with aiosqlite.connect(db_path) as db:
        # Only one active survey at a time — the currently uploaded questionnaire
        # is what clients see; older surveys stay in the DB (their responses and
        # exports remain intact) but stop accepting new answers.
        await db.execute("UPDATE surveys SET is_active = 0 WHERE is_active = 1")
        cursor = await db.execute("INSERT INTO surveys (title, is_active) VALUES (?, 1)", (title,))
        survey_id = cursor.lastrowid
        await db.executemany(
            "INSERT INTO survey_questions (survey_id, position, question) VALUES (?, ?, ?)",
            [(survey_id, i, q) for i, q in enumerate(questions)],
        )
        await db.commit()
        return survey_id


async def _get_active_survey(db_path: str) -> dict | None:
    async with aiosqlite.connect(db_path) as db:
        db.row_factory = aiosqlite.Row
        survey = await (await db.execute(
            "SELECT id, title FROM surveys WHERE is_active = 1 ORDER BY id DESC LIMIT 1"
        )).fetchone()
        if not survey:
            return None
        questions = await (await db.execute(
            "SELECT id, question FROM survey_questions WHERE survey_id = ? ORDER BY position",
            (survey["id"],),
        )).fetchall()
        return {
            "id": survey["id"],
            "title": survey["title"],
            "questions": [{"id": q["id"], "question": q["question"]} for q in questions],
        }


async def _save_response(
    db_path: str, survey_id: int, user_telegram_id: int, user_name: str | None,
    answers: list[tuple[int, str]],
) -> int:
    async with aiosqlite.connect(db_path) as db:
        cursor = await db.execute(
            "INSERT INTO survey_responses (survey_id, user_telegram_id, user_name) VALUES (?, ?, ?)",
            (survey_id, user_telegram_id, user_name),
        )
        response_id = cursor.lastrowid
        await db.executemany(
            "INSERT INTO survey_answers (response_id, question_id, answer) VALUES (?, ?, ?)",
            [(response_id, qid, answer) for qid, answer in answers],
        )
        await db.commit()
        return response_id


async def _response_count(db_path: str, survey_id: int) -> int:
    async with aiosqlite.connect(db_path) as db:
        row = await (await db.execute(
            "SELECT COUNT(*) FROM survey_responses WHERE survey_id = ?", (survey_id,)
        )).fetchone()
        return row[0] if row else 0


async def _fetch_export_rows(db_path: str, survey_id: int) -> tuple[list[str], list[dict]]:
    """Returns (question texts in order, rows) where each row is
    {"user_name", "submitted_at", "answers": [text ordered like questions]} —
    shared by both the table/preview text and the Word/Excel export so the
    column order can never drift between the two."""
    async with aiosqlite.connect(db_path) as db:
        db.row_factory = aiosqlite.Row
        questions = await (await db.execute(
            "SELECT id, question FROM survey_questions WHERE survey_id = ? ORDER BY position",
            (survey_id,),
        )).fetchall()
        q_ids = [q["id"] for q in questions]
        q_texts = [q["question"] for q in questions]

        responses = await (await db.execute(
            "SELECT id, user_telegram_id, user_name, submitted_at FROM survey_responses "
            "WHERE survey_id = ? ORDER BY id",
            (survey_id,),
        )).fetchall()

        rows = []
        for resp in responses:
            answer_rows = await (await db.execute(
                "SELECT question_id, answer FROM survey_answers WHERE response_id = ?",
                (resp["id"],),
            )).fetchall()
            by_qid = {r["question_id"]: r["answer"] for r in answer_rows}
            rows.append({
                "user_name": resp["user_name"] or str(resp["user_telegram_id"]),
                "submitted_at": resp["submitted_at"],
                "answers": [by_qid.get(qid, "") for qid in q_ids],
            })
        return q_texts, rows


# ── file/voice parsing ───────────────────────────────────────────────────────

async def _extract_text_from_document(bot: Bot, document: Document) -> str:
    """Downloads a PDF/Word/Excel document and extracts its raw text.
    Sizes are pre-checked by the caller (FILE_MAX_SIZE_BYTES) before this
    ever runs, so no additional bound here."""
    tg_file = await bot.get_file(document.file_id)
    buf = io.BytesIO()
    await bot.download_file(tg_file.file_path, destination=buf)
    buf.seek(0)
    name = (document.file_name or "").lower()

    if name.endswith(".pdf"):
        reader = PdfReader(buf)
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    if name.endswith(".docx"):
        from docx import Document as DocxDocument
        doc = DocxDocument(buf)
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                lines.extend(cell.text for cell in row.cells if cell.text.strip())
        return "\n".join(lines)

    if name.endswith(".xlsx"):
        from openpyxl import load_workbook
        wb = load_workbook(buf, read_only=True, data_only=True)
        lines = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                lines.extend(str(cell) for cell in row if cell is not None and str(cell).strip())
        return "\n".join(lines)

    return ""


async def _transcribe_voice(bot: Bot, voice) -> str:
    """Same AssemblyAI upload/poll pattern as features/voice_intake.py's
    _transcribe_voice — duplicated rather than imported because that module's
    function is private and tied to its own VoiceSchema-based dispatch, which
    this template does not use (it parses free-form questionnaire text, not a
    typed record)."""
    tg_file = await bot.get_file(voice.file_id)
    audio_url = f"https://api.telegram.org/file/bot{bot.token}/{tg_file.file_path}"
    async with aiohttp.ClientSession() as s:
        async with s.get(audio_url) as r:
            audio_bytes = await r.read()
        async with s.post(
            "https://api.assemblyai.com/v2/upload",
            headers={"authorization": ASSEMBLYAI_API_KEY, "content-type": "audio/ogg"},
            data=audio_bytes,
        ) as r:
            up = await r.json()
        upload_url = up.get("upload_url", "")
        if not upload_url:
            return ""
        async with s.post(
            "https://api.assemblyai.com/v2/transcript",
            headers={"authorization": ASSEMBLYAI_API_KEY, "content-type": "application/json"},
            json={"audio_url": upload_url, "language_code": VOICE_LANGUAGE_CODE},
        ) as r:
            tx = await r.json()
        tx_id = tx.get("id", "")
        if not tx_id:
            return ""
        for _ in range(60):
            await asyncio.sleep(3)
            async with s.get(
                f"https://api.assemblyai.com/v2/transcript/{tx_id}",
                headers={"authorization": ASSEMBLYAI_API_KEY},
            ) as r:
                res = await r.json()
            if res.get("status") == "completed":
                return res.get("text", "")
            if res.get("status") == "error":
                return ""
    return ""


def _split_questions_locally(text: str) -> list[str]:
    """Cheap heuristic fallback used both when Claude parsing fails/is
    unavailable AND as the primary path for plain typed text (no need to pay
    an API call for "each line is a question") — strips leading numbering
    ("1.", "1)", "- ") from each non-empty line."""
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    cleaned = [re.sub(r"^(\d+[\.\)]|[-•*])\s*", "", ln).strip() for ln in lines]
    return [c for c in cleaned if c][:MAX_QUESTIONS]


async def _parse_questions_with_claude(text: str) -> list[str] | None:
    """Used for voice transcripts and file extracts, where raw text rarely
    lines up one-question-per-line the way typed input does. Returns None on
    any failure so the caller can fall back to the local heuristic instead of
    silently producing an empty questionnaire."""
    prompt = (
        "Ты помощник, который выделяет вопросы анкеты из произвольного текста "
        "(может быть с распознавания голоса или из документа). Верни JSON-массив "
        "строк — только сами вопросы, без нумерации и лишних пояснений.\n\n"
        f"Текст:\n{text[:8000]}\n\n"
        "Ответ ТОЛЬКО JSON-массивом строк, например [\"Вопрос 1?\", \"Вопрос 2?\"]"
    )
    try:
        async with aiohttp.ClientSession() as s:
            async with s.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key": ANTHROPIC_API_KEY,
                    "anthropic-version": "2023-06-01",
                    "content-type": "application/json",
                },
                json={
                    "model": "claude-haiku-4-5-20251001",
                    "max_tokens": 1024,
                    "messages": [{"role": "user", "content": prompt}],
                },
            ) as r:
                res = await r.json()
        raw = res["content"][0]["text"]
        raw = re.sub(r"^```(?:json)?\n?", "", raw.strip())
        raw = re.sub(r"\n?```$", "", raw.strip())
        parsed = json.loads(raw)
        if not isinstance(parsed, list):
            return None
        questions = [str(q).strip()[:QUESTION_MAX_LEN] for q in parsed if str(q).strip()]
        return questions[:MAX_QUESTIONS] or None
    except Exception:
        logger.exception("survey_form: Claude question parsing failed")
        return None


# ── export: table preview / Excel ───────────────────────────────────────────

def _build_table_preview_text(title: str, q_texts: list[str], rows: list[dict]) -> str:
    lines = [f"📊 <b>{_esc(title)}</b> — {len(rows)} ответ(ов)\n"]
    if not rows:
        lines.append("Ответов пока нет.")
        return _join_bounded(lines)
    for row in rows[:RECENT_RESPONSES_LIMIT]:
        lines.append(f"<b>{_esc(row['user_name'])}</b> · {row['submitted_at']}")
        for q, a in zip(q_texts, row["answers"]):
            lines.append(f"  • {_esc(q, 200)}: {_esc(a, ANSWER_MAX_LEN)}")
        lines.append("")
    if len(rows) > RECENT_RESPONSES_LIMIT:
        lines.append(f"… и ещё {len(rows) - RECENT_RESPONSES_LIMIT}. Полный список — в экспорте.")
    return _join_bounded(lines)


def _build_excel_bytes(title: str, q_texts: list[str], rows: list[dict]) -> bytes:
    """Word-wrapped columns so long answers never make the sheet look
    truncated/misaligned — matches the ТЗ's "перенос текста без съезжания"
    requirement."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Ответы"[:31]

    headers = ["Имя", "Дата"] + q_texts
    ws.append(headers)
    for col_idx in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.font = Font(bold=True)
        cell.alignment = Alignment(wrap_text=True, vertical="top")

    for row in rows:
        ws.append([row["user_name"], row["submitted_at"], *row["answers"]])

    for row_cells in ws.iter_rows(min_row=2):
        for cell in row_cells:
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    for col_idx, header in enumerate(headers, start=1):
        ws.column_dimensions[get_column_letter(col_idx)].width = min(40, max(15, len(str(header)) + 5))

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


# ── FSM states ───────────────────────────────────────────────────────────────

class UploadFlow(StatesGroup):
    awaiting_input = State()
    confirm = State()

class AnswerFlow(StatesGroup):
    answering = State()

class AdminMgmtFlow(StatesGroup):
    add_admin = State()
    remove_admin_pick = State()


# ── keyboards ─────────────────────────────────────────────────────────────────

def kb_back(callback_data: str = "main_menu") -> InlineKeyboardButton:
    return InlineKeyboardButton(text="◀️ Назад", callback_data=callback_data)

def kb_flow_cancel() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="❌ Отмена", callback_data="sv_flow_cancel")],
    ])

def kb_main_menu() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📥 Загрузить анкету", callback_data="sv_upload")],
        [InlineKeyboardButton(text="📊 Ответы", callback_data="sv_responses")],
        [InlineKeyboardButton(text="👥 Админы", callback_data="adm_menu")],
    ])

def kb_confirm_survey() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✅ Всё верно", callback_data="sv_confirm_yes")],
        [InlineKeyboardButton(text="❌ Отмена", callback_data="sv_flow_cancel")],
    ])

def kb_responses_menu() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📄 Экспорт в Excel", callback_data="sv_export_xlsx")],
        [InlineKeyboardButton(text="📝 Экспорт в Word", callback_data="sv_export_docx")],
        [kb_back()],
    ])

def kb_admins_menu() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="➕ Добавить админа", callback_data="adm_add")],
        [InlineKeyboardButton(text="➖ Убрать админа", callback_data="adm_remove")],
        [kb_back()],
    ])

MAX_ADMIN_REMOVE_BUTTONS = 30

def kb_remove_admins(ids: list[str]) -> InlineKeyboardMarkup:
    rows = [[InlineKeyboardButton(text=admin_id, callback_data=f"adm_rm:{i}")] for i, admin_id in enumerate(ids)]
    rows.append([kb_back("adm_menu")])
    return InlineKeyboardMarkup(inline_keyboard=rows)


# ── /start ────────────────────────────────────────────────────────────────────

@router.message(Command("start"), F.chat.type == "private")
async def cmd_start(message: Message, state: FSMContext, config: SurveyFormConfig):
    await state.clear()
    admins = _load_admins(config.admins_file)
    first_time_admin = not admins
    if first_time_admin:
        _save_admins(config.admins_file, {str(message.from_user.id)})
        admins = {str(message.from_user.id)}

    if str(message.from_user.id) in admins:
        if config.welcome_image.exists():
            await message.answer_photo(
                FSInputFile(str(config.welcome_image)), caption=WELCOME_TEXT,
                parse_mode="HTML", reply_markup=kb_main_menu(),
            )
        else:
            await message.answer(WELCOME_TEXT, parse_mode="HTML", reply_markup=kb_main_menu())
        if first_time_admin:
            await message.answer(
                "👑 <b>Вы — администратор этого бота.</b>\n\n"
                "Управление другими администраторами — кнопка «👥 Админы» выше.",
                parse_mode="HTML",
            )
        return

    survey = await _get_active_survey(config.db_path)
    if not survey or not survey["questions"]:
        await message.answer("Сейчас анкета не заполняется, загляните позже 🙂")
        return

    await state.set_state(AnswerFlow.answering)
    await state.update_data(
        survey_id=survey["id"], questions=survey["questions"], idx=0, answers=[], started_at=time.time(),
    )
    await message.answer(CLIENT_WELCOME_TEXT, parse_mode="HTML")
    await message.answer(_esc(survey["questions"][0]["question"], QUESTION_MAX_LEN))


@router.callback_query(F.data == "main_menu")
async def cb_main_menu(cb: CallbackQuery, state: FSMContext, config: SurveyFormConfig):
    await cb.answer()
    if not _is_admin(cb.from_user.id, config):
        return
    await state.clear()
    await cb.message.edit_text(WELCOME_TEXT, parse_mode="HTML", reply_markup=kb_main_menu())


@router.callback_query(F.data == "sv_flow_cancel")
async def cb_flow_cancel(cb: CallbackQuery, state: FSMContext, config: SurveyFormConfig):
    await cb.answer()
    await state.clear()
    if _is_admin(cb.from_user.id, config):
        await cb.message.edit_text(WELCOME_TEXT, parse_mode="HTML", reply_markup=kb_main_menu())
    else:
        await cb.message.edit_text("Хорошо, может быть в другой раз! 🙂")


# ── admin: upload survey ─────────────────────────────────────────────────────

@router.callback_query(F.data == "sv_upload")
async def cb_upload_start(cb: CallbackQuery, state: FSMContext, config: SurveyFormConfig):
    await cb.answer()
    if not _is_admin(cb.from_user.id, config):
        return
    await state.set_state(UploadFlow.awaiting_input)
    await state.update_data(started_at=time.time())
    await cb.message.edit_text(UPLOAD_PROMPT_TEXT, parse_mode="HTML", reply_markup=kb_flow_cancel())


def _questions_confirm_text(questions: list[str]) -> str:
    lines = [f"Вот {len(questions)} вопрос(ов). Правильно?\n"]
    lines.extend(f"{i}. {_esc(q, QUESTION_MAX_LEN)}" for i, q in enumerate(questions, 1))
    return _join_bounded(lines)


async def _present_parsed_questions(message: Message, state: FSMContext, questions: list[str]) -> None:
    if not questions:
        await message.answer(
            "Не удалось распознать вопросы. Попробуйте прислать текст, голос или файл ещё раз.",
            reply_markup=kb_flow_cancel(),
        )
        return
    await state.update_data(parsed_questions=questions)
    await state.set_state(UploadFlow.confirm)
    await message.answer(_questions_confirm_text(questions), parse_mode="HTML", reply_markup=kb_confirm_survey())


@router.message(UploadFlow.awaiting_input, F.text, ~F.text.startswith("/"))
async def upload_text(msg: Message, state: FSMContext):
    data = await state.get_data()
    if _flow_expired(data):
        await state.clear()
        await msg.answer("Сессия истекла, начните заново.")
        return
    questions = _split_questions_locally(msg.text)
    await _present_parsed_questions(msg, state, questions)


@router.message(UploadFlow.awaiting_input, F.voice)
async def upload_voice(msg: Message, state: FSMContext, bot: Bot):
    data = await state.get_data()
    if _flow_expired(data):
        await state.clear()
        await msg.answer("Сессия истекла, начните заново.")
        return
    status = await msg.answer("🎙️ Распознаю голосовое сообщение…")
    text = await _transcribe_voice(bot, msg.voice)
    if not text:
        await status.edit_text("Не удалось распознать голосовое сообщение. Попробуйте ещё раз.")
        return
    questions = await _parse_questions_with_claude(text) or _split_questions_locally(text)
    await status.delete()
    await _present_parsed_questions(msg, state, questions)


@router.message(UploadFlow.awaiting_input, F.document)
async def upload_document(msg: Message, state: FSMContext, bot: Bot):
    data = await state.get_data()
    if _flow_expired(data):
        await state.clear()
        await msg.answer("Сессия истекла, начните заново.")
        return
    doc = msg.document
    name = (doc.file_name or "").lower()
    if not name.endswith((".pdf", ".docx", ".xlsx")):
        await msg.answer("Поддерживаются файлы PDF, Word (.docx) и Excel (.xlsx).")
        return
    if doc.file_size and doc.file_size > FILE_MAX_SIZE_BYTES:
        await msg.answer(f"Файл слишком большой (максимум {FILE_MAX_SIZE_BYTES // (1024 * 1024)} МБ).")
        return

    status = await msg.answer("📄 Разбираю файл…")
    try:
        text = await _extract_text_from_document(bot, doc)
    except Exception:
        logger.exception("survey_form: failed to extract text from uploaded document")
        text = ""
    if not text.strip():
        await status.edit_text("Не удалось прочитать файл. Попробуйте другой формат или пришлите текст.")
        return
    questions = await _parse_questions_with_claude(text) or _split_questions_locally(text)
    await status.delete()
    await _present_parsed_questions(msg, state, questions)


@router.callback_query(UploadFlow.confirm, F.data == "sv_confirm_yes")
async def cb_confirm_survey(cb: CallbackQuery, state: FSMContext, config: SurveyFormConfig):
    await cb.answer()
    data = await state.get_data()
    if _flow_expired(data):
        await state.clear()
        await cb.message.edit_text("Сессия истекла, начните заново.")
        return
    questions = data.get("parsed_questions") or []
    if not questions:
        await state.clear()
        await cb.message.edit_text("Список вопросов пуст.", reply_markup=kb_main_menu())
        return
    title = f"Анкета от {time.strftime('%Y-%m-%d')}"
    survey_id = await _create_survey(config.db_path, title, questions)
    await state.clear()
    logger.info(f"survey_form: survey {survey_id} created by {cb.from_user.id} with {len(questions)} questions")
    await cb.message.edit_text(
        f"✅ Анкета сохранена и активна ({len(questions)} вопрос(ов)).", reply_markup=kb_main_menu(),
    )


# ── client: answering flow ───────────────────────────────────────────────────

async def _notify_admins_new_response(bot: Bot, config: SurveyFormConfig, survey_title: str, user_name: str) -> None:
    admins = _load_admins(config.admins_file)
    text = f"📋 Новый ответ на анкету «{_esc(survey_title)}» от {_esc(user_name)}."
    for admin_id in admins:
        try:
            await bot.send_message(int(admin_id), text, parse_mode="HTML")
        except Exception:
            logger.exception(f"survey_form: failed to notify admin {admin_id}")


@router.message(AnswerFlow.answering, F.text, ~F.text.startswith("/"))
async def answer_text(msg: Message, state: FSMContext, config: SurveyFormConfig, bot: Bot):
    data = await state.get_data()
    if _flow_expired(data):
        await state.clear()
        await msg.answer("Сессия истекла, начните заново — отправьте /start.")
        return

    questions = data["questions"]
    idx = data["idx"]
    answers = data["answers"]
    answers.append((questions[idx]["id"], msg.text.strip()[:ANSWER_MAX_LEN]))
    idx += 1

    if idx >= len(questions):
        survey_id = data["survey_id"]
        user_name = msg.from_user.full_name or msg.from_user.username or str(msg.from_user.id)
        await _save_response(config.db_path, survey_id, msg.from_user.id, user_name, answers)
        await state.clear()
        await msg.answer(THANK_YOU_TEXT)
        survey = await _get_active_survey(config.db_path)
        survey_title = survey["title"] if survey else "?"
        asyncio.create_task(_notify_admins_new_response(bot, config, survey_title, user_name))
        return

    await state.update_data(idx=idx, answers=answers)
    await msg.answer(_esc(questions[idx]["question"], QUESTION_MAX_LEN))


# ── admin: responses / export ────────────────────────────────────────────────

@router.callback_query(F.data == "sv_responses")
async def cb_responses_menu(cb: CallbackQuery, state: FSMContext, config: SurveyFormConfig):
    await cb.answer()
    if not _is_admin(cb.from_user.id, config):
        return
    await state.clear()
    survey = await _get_active_survey(config.db_path)
    if not survey:
        await cb.message.edit_text("Анкета ещё не загружена.", reply_markup=kb_main_menu())
        return
    q_texts, rows = await _fetch_export_rows(config.db_path, survey["id"])
    text = _build_table_preview_text(survey["title"], q_texts, rows)
    await cb.message.edit_text(text, parse_mode="HTML", reply_markup=kb_responses_menu())


@router.callback_query(F.data == "sv_export_xlsx")
async def cb_export_xlsx(cb: CallbackQuery, config: SurveyFormConfig):
    await cb.answer()
    if not _is_admin(cb.from_user.id, config):
        return
    survey = await _get_active_survey(config.db_path)
    if not survey:
        await cb.message.answer("Анкета ещё не загружена.")
        return
    q_texts, rows = await _fetch_export_rows(config.db_path, survey["id"])
    if not rows:
        await cb.message.answer("Ответов пока нет.")
        return
    data = _build_excel_bytes(survey["title"], q_texts, rows)
    filename = f"survey_{survey['id']}.xlsx"
    await cb.message.answer_document(BufferedInputFile(data, filename=filename))


@router.callback_query(F.data == "sv_export_docx")
async def cb_export_docx(cb: CallbackQuery, config: SurveyFormConfig):
    await cb.answer()
    if not _is_admin(cb.from_user.id, config):
        return
    survey = await _get_active_survey(config.db_path)
    if not survey:
        await cb.message.answer("Анкета ещё не загружена.")
        return
    q_texts, rows = await _fetch_export_rows(config.db_path, survey["id"])
    if not rows:
        await cb.message.answer("Ответов пока нет.")
        return
    for row in rows:
        fields = list(zip(q_texts, row["answers"]))
        data = build_document(
            f"{survey['title']} — {row['user_name']}", fields,
            bot_name=survey["title"], user_name=row["user_name"],
        )
        filename = f"survey_{survey['id']}_{row['user_name'][:30]}.docx"
        await cb.message.answer_document(BufferedInputFile(data, filename=filename))


# ── ADMINS menu ────────────────────────────────────────────────────────────────

async def _admins_list_text(config: SurveyFormConfig) -> str:
    ids = sorted(_load_admins(config.admins_file))
    if not ids:
        return "👥 Пусто"
    return _join_bounded(["👥 <b>Администраторы бота:</b>\n"] + [f"• <code>{_esc(i)}</code>" for i in ids])


@router.callback_query(F.data == "adm_menu")
async def cb_adm_menu(cb: CallbackQuery, state: FSMContext, config: SurveyFormConfig):
    await cb.answer()
    if not _is_admin(cb.from_user.id, config):
        return
    await state.clear()
    text = await _admins_list_text(config)
    await cb.message.edit_text(text, parse_mode="HTML", reply_markup=kb_admins_menu())


@router.callback_query(F.data == "adm_add")
async def cb_adm_add(cb: CallbackQuery, state: FSMContext, config: SurveyFormConfig):
    await cb.answer()
    if not _is_admin(cb.from_user.id, config):
        return
    await state.set_state(AdminMgmtFlow.add_admin)
    await state.update_data(started_at=time.time())
    await cb.message.edit_text("Введите Telegram ID нового администратора:", reply_markup=kb_flow_cancel())


@router.message(AdminMgmtFlow.add_admin, F.text, ~F.text.startswith("/"))
async def admin_add_id(msg: Message, state: FSMContext, config: SurveyFormConfig):
    data = await state.get_data()
    if _flow_expired(data):
        await state.clear()
        await msg.answer("Сессия истекла, начните заново.", reply_markup=kb_main_menu())
        return
    text = msg.text.strip()
    if not _valid_admin_id(text):
        await msg.answer("Некорректный ID. Введите числовой Telegram ID.", reply_markup=kb_flow_cancel())
        return
    await state.clear()
    ids = _load_admins(config.admins_file)
    ids.add(text)
    _save_admins(config.admins_file, ids)
    logger.info(f"survey_form: {text} added as bot admin by {msg.from_user.id}")
    await msg.answer(f"✅ <code>{text}</code> добавлен.", parse_mode="HTML", reply_markup=kb_admins_menu())


@router.callback_query(F.data == "adm_remove")
async def cb_adm_remove(cb: CallbackQuery, state: FSMContext, config: SurveyFormConfig):
    await cb.answer()
    if not _is_admin(cb.from_user.id, config):
        return
    ids = sorted(_load_admins(config.admins_file))
    if len(ids) <= 1:
        await cb.message.edit_text("Нельзя удалить последнего администратора.", reply_markup=kb_admins_menu())
        return
    if len(ids) > MAX_ADMIN_REMOVE_BUTTONS:
        await cb.message.edit_text(
            "Слишком много админов для списка кнопок. Обратитесь к разработчику.", reply_markup=kb_admins_menu()
        )
        return
    await state.set_state(AdminMgmtFlow.remove_admin_pick)
    await state.update_data(started_at=time.time(), remove_admin_ids=ids)
    await cb.message.edit_text("Выберите администратора для удаления:", reply_markup=kb_remove_admins(ids))


@router.callback_query(AdminMgmtFlow.remove_admin_pick, F.data.startswith("adm_rm:"))
async def cb_adm_remove_pick(cb: CallbackQuery, state: FSMContext, config: SurveyFormConfig):
    await cb.answer()
    if not _is_admin(cb.from_user.id, config):
        return
    data = await state.get_data()
    if _flow_expired(data):
        await state.clear()
        await cb.message.edit_text("Сессия истекла, начните заново.", reply_markup=kb_main_menu())
        return
    try:
        idx = int(cb.data.split(":", 1)[1])
        remove_admin_ids = data["remove_admin_ids"]
        if not (0 <= idx < len(remove_admin_ids)):
            raise IndexError(idx)
        target = remove_admin_ids[idx]
    except (ValueError, IndexError, KeyError):
        await state.clear()
        await cb.message.edit_text("Некорректный выбор.", reply_markup=kb_admins_menu())
        return
    ids = _load_admins(config.admins_file)
    if len(ids) <= 1:
        await state.clear()
        await cb.message.edit_text("Нельзя удалить последнего администратора.", reply_markup=kb_admins_menu())
        return
    ids.discard(target)
    _save_admins(config.admins_file, ids)
    logger.info(f"survey_form: {target} removed as bot admin by {cb.from_user.id}")
    await state.clear()
    await cb.message.edit_text(f"✅ <code>{_esc(target)}</code> удалён.", parse_mode="HTML",
                                reply_markup=kb_admins_menu())


# ── MAIN ──────────────────────────────────────────────────────────────────────

async def main():
    config = config_from_env()
    bot = Bot(token=os.getenv("BOT_TOKEN"))
    dp = Dispatcher(storage=MemoryStorage())
    dp.update.outer_middleware(ConfigMiddleware(config))
    dp.include_router(router)
    await bot.set_my_description(BOT_DESCRIPTION)
    await init_db(config.db_path)
    await dp.start_polling(bot)

if __name__ == "__main__":
    asyncio.run(main())
