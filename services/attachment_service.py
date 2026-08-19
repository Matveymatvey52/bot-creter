"""Turns Telegram photos/documents into Claude Messages API content blocks
for the /create requirements-gathering dialog (handlers/create_bot.py).

Content blocks are kept as plain JSON-serializable dicts (base64 strings,
plain text) rather than Python objects, so the same buffer format will also
fit stage D (Google Sheets share-link text) without any rework — a link
just becomes another {"type": "text", ...} block appended to the buffer.
"""
from __future__ import annotations

import io
import logging

logger = logging.getLogger(__name__)

MAX_DOCUMENT_SIZE_BYTES = 15 * 1024 * 1024
MAX_EXTRACTED_CHARS = 20_000
MAX_PENDING_ATTACHMENTS = 6

SUPPORTED_IMAGE_MIME = {"image/jpeg", "image/png", "image/gif", "image/webp"}


def build_image_block(data: bytes, media_type: str) -> dict:
    import base64

    return {
        "type": "image",
        "source": {
            "type": "base64",
            "media_type": media_type,
            "data": base64.b64encode(data).decode("ascii"),
        },
    }


def build_text_block(text: str) -> dict:
    return {"type": "text", "text": text}


def _truncate(text: str) -> str:
    text = text.strip()
    if len(text) > MAX_EXTRACTED_CHARS:
        text = text[:MAX_EXTRACTED_CHARS] + "\n…[обрезано]"
    return text


def extract_pdf_text(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = [page.extract_text() or "" for page in reader.pages]
    return _truncate("\n".join(parts))


def extract_docx_text(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return _truncate("\n".join(parts))


def extract_xlsx_text(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"[Лист: {sheet.title}]")
        for i, row in enumerate(sheet.iter_rows(values_only=True)):
            if i >= 200:
                parts.append("…")
                break
            cells = ["" if c is None else str(c) for c in row]
            if any(c.strip() for c in cells):
                parts.append(" | ".join(cells))
    return _truncate("\n".join(parts))


# (extension, extractor) — checked in order against the lowercased filename
_EXTRACTORS = (
    (".pdf", extract_pdf_text),
    (".docx", extract_docx_text),
    (".xlsx", extract_xlsx_text),
    (".xls", extract_xlsx_text),
)


def extract_document_text(data: bytes, filename: str) -> str | None:
    """Returns extracted text, or None if the file type isn't supported."""
    name = (filename or "").lower()
    for ext, extractor in _EXTRACTORS:
        if name.endswith(ext):
            try:
                return extractor(data)
            except Exception:
                logger.exception("Failed to extract text from %s", filename)
                raise
    return None
