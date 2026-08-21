# FEATURE: channel_monitor
# COMPATIBLE_WITH: channel_aggregator
"""Reusable Telegram channel-monitoring feature module. As of the owner's
2026-08-21 decision, channel monitoring is offered to factory users ONLY as
its own standalone bot type (templates/channel_aggregator.py, which imports
this module's router/init_db directly and auto-enables it — see that file's
docstring) — this module is no longer offered as a pluggable feature for
arbitrary templates (COMPATIBLE_WITH used to be "*", letting the owner attach
it to any bot; that is reverted here). COMPATIBLE_WITH now names only
channel_aggregator, so handlers/manage_bots.py's Features panel never offers
it to a NEW bot of any other template.

features/channel_monitor.py itself is deliberately NOT deleted and its module
path is unchanged: bot_id=12 (tour_operator_demo) already has "channel_monitor"
enabled in bot_features from before this decision, and
runtime/registry.py's _load_and_include_features() imports feature modules by
that exact bot_features name — deleting or renaming the file would break that
bot, which the owner explicitly asked to leave untouched. It keeps working,
"orphaned": functional for bot 12, invisible to everyone else's Features
panel. See docs/USERBOT_FEATURE_DESIGN.md for the (now superseded) design
that made this a universal "*" feature, and
docs/CHANNEL_AGGREGATOR_TEMPLATE_DESIGN.md for the standalone-template design
that replaces it as the only way to get channel monitoring on a new bot.

Exposes a module-level `router` so runtime/registry.py's
_load_and_include_features() can clone and attach it to bot_id=12 (the sole
remaining bot with "channel_monitor" enabled in bot_features) and
templates/channel_aggregator.py can wire it directly into its own Dispatcher.
This module does NOT own the host bot's /start — it only adds a "📡 Мониторинг
каналов" entry point (button/command) into whatever menu the host template
already has, via `bot_commands` below.

Session ownership (design doc §2, Variant A): one userbot Telethon session
belongs to exactly ONE bot_id, stored in that bot's OWN per-bot db_path (same
"features reuse the host template's own per-bot db_path" convention as every
other feature) — never a factory-wide pool shared across bots. Revoking
access for one bot never touches any other bot's session.

The actual long-lived Telethon connection and post ingestion live in a
SEPARATE process, runtime/userbot_worker.py — this module only talks to
Telegram's Bot API (aiogram) plus the per-bot db_path tables both processes
share (db/database.py's userbot_sessions / monitored_channels / channel_posts,
now bot_id-scoped).
"""
from __future__ import annotations

import csv
import html
import io
import json
import logging
import re
from datetime import datetime, timedelta
from pathlib import Path
from typing import Protocol

from aiogram import F, Router
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.types import (
    BufferedInputFile, CallbackQuery, InlineKeyboardButton, InlineKeyboardMarkup, Message,
)

from config import TELEGRAM_API_HASH, TELEGRAM_API_ID
from db.database import (
    add_monitored_channel,
    count_monitored_channels,
    create_userbot_session,
    get_monitored_channels_by_bot,
    get_posts_for_report,
    get_recent_posts_for_bot,
    get_userbot_sessions_by_bot,
    init_channel_monitor_tables,
    revoke_userbot_session,
    set_monitored_channel_active,
    set_monitored_channel_schema,
    set_report_schedule,
    set_userbot_session_active,
    set_userbot_session_status,
)

logger = logging.getLogger(__name__)
router = Router()

bot_commands = [("monitor", "Мониторинг каналов")]

MAX_LIST_BUTTONS = 30
FEED_LIMIT = 15
TEXT_PREVIEW_LEN = 400
MAX_CHANNELS_PER_BOT = 50  # owner decision #4, docs/CHANNEL_AGGREGATOR_TEMPLATE_DESIGN.md §5

# ── schema presets (owner decision #2: wide, extensible list, not just 4) ────
# Each preset is {key, label, fields}. "Свободная схема" (custom) and
# "Без схемы" (none) are handled separately in the picker, not in this list —
# they aren't a fixed field set. Extending this list is the only change
# needed to add a new niche preset; nothing else in the FSM hardcodes them.
SCHEMA_PRESETS: list[dict] = [
    {"key": "jobs", "label": "💼 Вакансии", "fields": ["должность", "зарплата", "локация", "контакт"]},
    {"key": "news", "label": "📰 Новости/анонсы", "fields": ["заголовок", "дата", "ключевые темы", "суть в 1 фразе"]},
    {"key": "classifieds", "label": "🏷 Объявления/лоты", "fields": ["товар", "цена", "контакт"]},
    {"key": "engagement", "label": "📈 Аналитика/охваты", "fields": ["заголовок", "дата", "ключевые темы", "метрики вовлечённости"]},
    {"key": "realestate", "label": "🏠 Недвижимость", "fields": ["тип объекта", "цена", "площадь", "район", "контакт"]},
    {"key": "ecommerce", "label": "🛒 E-commerce/цены", "fields": ["товар", "цена", "скидка", "магазин", "ссылка"]},
    {"key": "events", "label": "🎫 Мероприятия/анонсы", "fields": ["название события", "дата", "место", "цена билета"]},
    {"key": "reviews", "label": "⭐ Обзоры/отзывы", "fields": ["объект обзора", "оценка", "плюсы", "минусы"]},
    {"key": "finance", "label": "💹 Финансы/крипто-сигналы", "fields": ["актив", "направление", "цена входа", "цель", "стоп-лосс"]},
    {"key": "hr", "label": "🧑‍💼 HR/резюме", "fields": ["имя", "специализация", "опыт", "контакт"]},
    {"key": "competitors", "label": "🕵️ Конкуренты/маркетинг", "fields": ["компания", "продукт", "акция", "дата"]},
    {"key": "tech", "label": "💻 IT/технологии", "fields": ["тема", "продукт/технология", "суть в 1 фразе"]},
    {"key": "tourism", "label": "✈️ Туризм", "fields": ["направление", "цена", "даты", "оператор"]},
]
_PRESETS_BY_KEY = {p["key"]: p for p in SCHEMA_PRESETS}

REPORT_PERIODS = {
    "today": "Сегодня",
    "week": "Неделя",
    "all": "Всё время",
}

WEEKDAY_NAMES = ["понедельник", "вторник", "среда", "четверг", "пятница", "суббота", "воскресенье"]

CLOSED_CHANNEL_GUIDE_TEXT = (
    "❓ <b>Как добавить закрытый канал</b>\n\n"
    "Мониторинг работает через ваш личный Telegram-аккаунт — он должен САМ "
    "состоять в канале, чтобы читать его посты. Для <b>публичных</b> каналов "
    "это происходит автоматически по @username.\n\n"
    "Для <b>закрытого</b> (приватного) канала есть два способа:\n"
    "1. Попросите администратора канала добавить ваш аккаунт вручную (по "
    "номеру телефона или username) — после этого канал можно подключить как "
    "обычно, по @username.\n"
    "2. Получите инвайт-ссылку канала у администратора и перешлите/сообщите "
    "её вашему аккаунту вручную, вступив в канал самостоятельно, — после "
    "этого он тоже станет доступен для подключения.\n\n"
    "Автоматическое вступление по инвайт-ссылке через бота сейчас не "
    "поддерживается — это осознанное решение, чтобы не вступать в чужие "
    "закрытые каналы без явного участия владельца аккаунта."
)

RISK_TEXT = (
    "⚠️ Мониторинг чужих каналов работает через ваш <b>личный Telegram-аккаунт</b>, "
    "а не через бота. Это значит:\n\n"
    "• Telegram может ограничить или заблокировать этот аккаунт, если сочтёт его "
    "поведение автоматизированным (массовое вступление в каналы, частые запросы) — "
    "такой риск реален и неустраним полностью на нашей стороне.\n"
    "• Массовый сбор контента из чужих каналов может нарушать условия использования "
    "Telegram — ответственность за то, какие каналы вы мониторите и как используете "
    "собранные данные, лежит на вас.\n"
    "• Мы технически предоставляем инструмент для авторизации и мониторинга, но не "
    "гарантируем сохранность вашего аккаунта и не несём ответственности за его "
    "блокировку.\n\n"
    "Продолжая, вы подтверждаете, что понимаете эти риски."
)

WELCOME_TEXT = (
    "📡 <b>Мониторинг каналов</b>\n\n"
    "Слежу за чужими Telegram-каналами через ваш личный аккаунт и собираю "
    "их посты в одну ленту — с кратким summary через Gemini.\n\n"
    "Нажмите «➕ Подключить мониторинг», чтобы начать."
)

MAX_CODE_ATTEMPTS = 3
MAX_2FA_ATTEMPTS = 3

# Security-review finding (carried over from templates/channel_monitor.py):
# without this, a Telegram user could type a different phone number into the
# bot on every /monitor -> cm_connect cycle, making this bot send a real
# Telegram login code to an arbitrary phone number over and over — an
# SMS/code-bombing proxy against numbers the requester doesn't own.
# Telegram's own FloodWaitError eventually kicks in, but only per phone
# number / per app, not per REQUESTING Telegram user — so it doesn't stop one
# user from cycling through many different victim numbers. This is a
# per-(bot_id, from_user.id) cooldown — keyed by bot_id too since the SAME
# feature module instance now serves many different bots (module-level dict,
# unlike per-bot instance state), independent of and in addition to
# FloodWaitError handling below.
PHONE_REQUEST_COOLDOWN_SECONDS = 30
_last_phone_request_at: dict[tuple[int, int], float] = {}


class ChannelMonitorConfig(Protocol):
    """Any template Config dataclass with db_path + admins_file works —
    duck-typed so this module doesn't need to know about any specific
    template's Config shape, same convention as features/payments.py's
    PaymentsConfig. admins_file added alongside the SECURITY fix below (see
    _is_bot_admin) — same {"ids": [...]} JSON shape every template's
    admins_file already has (features/sellable_items.py verified this is
    uniform across all COMPATIBLE_WITH templates)."""
    db_path: str
    admins_file: Path | str


# _load_admins/_is_bot_admin duplicated on purpose — no shared
# features/_common.py exists yet in this project (see
# features/sellable_items.py's own copy of the same helpers).
def _load_admins(admins_file: Path | str) -> set[str]:
    try:
        return set(json.loads(Path(admins_file).read_text()).get("ids", []))
    except Exception as e:
        logger.warning(f"_load_admins: failed to read {admins_file!r}: {e}")
        return set()


def _is_bot_admin(user_id: int, config: ChannelMonitorConfig) -> bool:
    return str(user_id) in _load_admins(config.admins_file)


async def init_db(db_path: str) -> None:
    await init_channel_monitor_tables(db_path)


# ── FSM ──────────────────────────────────────────────────────────────────────

class UserbotAuthStates(StatesGroup):
    waiting_risk_ack = State()
    waiting_phone = State()
    waiting_code = State()
    waiting_2fa_password = State()


class _AddChannelState(StatesGroup):
    waiting_username = State()
    waiting_schema_choice = State()
    waiting_custom_fields = State()


class _ScheduleState(StatesGroup):
    waiting_frequency = State()
    waiting_time = State()
    waiting_weekday = State()


class _ReportState(StatesGroup):
    waiting_channel_choice = State()


# ── keyboards ────────────────────────────────────────────────────────────────

def kb_start_menu() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="➕ Подключить мониторинг", callback_data="cm_connect")],
        [InlineKeyboardButton(text="📋 Мои каналы", callback_data="cm_channels")],
        [InlineKeyboardButton(text="📰 Лента", callback_data="cm_feed")],
        [InlineKeyboardButton(text="📊 Отчёт", callback_data="cm_report")],
        [InlineKeyboardButton(text="⏰ Настроить расписание", callback_data="cm_schedule")],
        [InlineKeyboardButton(text="❓ Как добавить закрытый канал", callback_data="cm_closed_guide")],
    ])


def kb_risk_ack() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✅ Понимаю и продолжаю", callback_data="cm_risk_ack")],
        [InlineKeyboardButton(text="❌ Отмена", callback_data="cm_cancel")],
    ])


def kb_cancel() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[[
        InlineKeyboardButton(text="❌ Отмена", callback_data="cm_cancel")
    ]])


def kb_channels(channels: list[dict]) -> InlineKeyboardMarkup:
    rows = []
    for ch in channels[:MAX_LIST_BUTTONS]:
        label = ch.get("channel_title") or ch.get("channel_username") or f"#{ch['id']}"
        state_icon = "🟢" if ch.get("active") else "⚪️"
        label = f"{state_icon} {label}"
        rows.append([InlineKeyboardButton(text=label[:60], callback_data=f"cm_toggle:{ch['id']}")])
    rows.append([InlineKeyboardButton(text="➕ Добавить канал", callback_data="cm_add_channel")])
    rows.append([InlineKeyboardButton(text="⬅️ Назад", callback_data="cm_menu")])
    return InlineKeyboardMarkup(inline_keyboard=rows)


def kb_back_to_menu() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[[
        InlineKeyboardButton(text="⬅️ Назад", callback_data="cm_menu")
    ]])


def kb_schema_picker() -> InlineKeyboardMarkup:
    rows = [[InlineKeyboardButton(text=p["label"], callback_data=f"cm_schema:{p['key']}")] for p in SCHEMA_PRESETS]
    rows.append([InlineKeyboardButton(text="✏️ Свободная схема", callback_data="cm_schema:custom")])
    rows.append([InlineKeyboardButton(text="🚫 Без схемы", callback_data="cm_schema:none")])
    # state-db review finding: every other keyboard in this file has an escape
    # hatch (cancel/back) — this one didn't, leaving the owner with no button
    # to leave the screen without picking something.
    rows.append([InlineKeyboardButton(text="❌ Отмена", callback_data="cm_cancel")])
    return InlineKeyboardMarkup(inline_keyboard=rows)


def kb_report_periods() -> InlineKeyboardMarkup:
    rows = [[InlineKeyboardButton(text=label, callback_data=f"cm_report_period:{key}")]
            for key, label in REPORT_PERIODS.items()]
    rows.append([InlineKeyboardButton(text="⬅️ Назад", callback_data="cm_menu")])
    return InlineKeyboardMarkup(inline_keyboard=rows)


def kb_report_channels(channels: list[dict]) -> InlineKeyboardMarkup:
    rows = [[InlineKeyboardButton(text="Все каналы", callback_data="cm_report_channel:all")]]
    for ch in channels[:MAX_LIST_BUTTONS]:
        label = ch.get("channel_title") or ch.get("channel_username") or f"#{ch['id']}"
        rows.append([InlineKeyboardButton(text=label[:60], callback_data=f"cm_report_channel:{ch['id']}")])
    rows.append([InlineKeyboardButton(text="⬅️ Назад", callback_data="cm_menu")])
    return InlineKeyboardMarkup(inline_keyboard=rows)


def kb_report_formats() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📄 CSV", callback_data="cm_report_fmt:csv")],
        [InlineKeyboardButton(text="📝 Word (.docx)", callback_data="cm_report_fmt:docx")],
        [InlineKeyboardButton(text="📊 Google Sheets", callback_data="cm_report_fmt:sheets")],
        [InlineKeyboardButton(text="⬅️ Назад", callback_data="cm_menu")],
    ])


def kb_schedule_frequency() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="Каждый день", callback_data="cm_freq:daily")],
        [InlineKeyboardButton(text="Раз в неделю", callback_data="cm_freq:weekly")],
        [InlineKeyboardButton(text="⬅️ Назад", callback_data="cm_menu")],
    ])


# ── helpers ──────────────────────────────────────────────────────────────────

def _esc(value, max_len: int = 500) -> str:
    text = str(value) if value is not None else ""
    if len(text) > max_len:
        text = text[:max_len] + "…"
    return html.escape(text)


def _make_client(session_string: str = ""):
    """Builds a fresh Telethon TelegramClient bound to a StringSession.

    Isolated into its own function (rather than inlined at each call site) so
    tests can patch features.channel_monitor._make_client instead of
    reaching into telethon.TelegramClient directly — no real network calls
    happen in tests. Raises a clear error if TELEGRAM_API_ID/API_HASH aren't
    configured, instead of Telethon's own more opaque failure deep inside
    connect().
    """
    if not TELEGRAM_API_ID or not TELEGRAM_API_HASH:
        raise RuntimeError(
            "TELEGRAM_API_ID / TELEGRAM_API_HASH is not set in .env — required to authorize "
            "a userbot session. Obtain both at https://my.telegram.org/apps and set them as "
            "environment variables."
        )
    from telethon import TelegramClient
    from telethon.sessions import StringSession

    return TelegramClient(StringSession(session_string), int(TELEGRAM_API_ID), TELEGRAM_API_HASH)


async def _channels_panel_text(db_path: str, bot_id: int) -> tuple[str, list[dict]]:
    channels = await get_monitored_channels_by_bot(db_path, bot_id)
    if not channels:
        return "У вас пока нет подключённых каналов.", channels
    lines = ["📋 <b>Ваши каналы</b> (нажмите, чтобы включить/выключить):\n"]
    for ch in channels[:MAX_LIST_BUTTONS]:
        icon = "🟢" if ch.get("active") else "⚪️"
        title = ch.get("channel_title") or ch.get("channel_username") or f"#{ch['id']}"
        lines.append(f"{icon} {_esc(title, 80)}")
    return "\n".join(lines), channels


async def _feed_text(db_path: str, bot_id: int) -> str:
    posts = await get_recent_posts_for_bot(db_path, bot_id, limit=FEED_LIMIT)
    if not posts:
        return "📰 Лента пуста — постов из активных каналов пока нет."
    lines = ["📰 <b>Лента</b>\n"]
    for p in posts:
        title = p.get("channel_title") or p.get("channel_username") or "канал"
        body = p.get("summary") or p.get("text") or ""
        lines.append(f"<b>{_esc(title, 60)}</b>\n{_esc(body, TEXT_PREVIEW_LEN)}\n")
    return "\n".join(lines)


# ── report generation (shared by "📊 Отчёт" handler AND the scheduled digest
# in runtime/userbot_worker.py — one implementation, no duplication) ─────────

def _post_row(post: dict) -> tuple[list[str], list[str]]:
    """Returns (headers, values) for one post — columns follow the channel's
    extract_schema fields (parsed out of `extracted` JSON) if set, otherwise a
    generic summary/text column, per design doc §3.4."""
    channel = post.get("channel_title") or post.get("channel_username") or ""
    schema_raw = post.get("extract_schema")
    fields: list[str] = []
    if schema_raw:
        try:
            fields = json.loads(schema_raw)
        except (ValueError, TypeError):
            fields = []
    if fields:
        extracted_raw = post.get("extracted")
        extracted: dict = {}
        if extracted_raw:
            try:
                extracted = json.loads(extracted_raw)
            except (ValueError, TypeError):
                extracted = {}
        headers = ["Канал", "Дата", *fields, "Просмотры", "Репосты"]
        values = [channel, str(post.get("posted_at") or ""), *[str(extracted.get(f, "")) for f in fields],
                  str(post.get("views") if post.get("views") is not None else ""),
                  str(post.get("forwards") if post.get("forwards") is not None else "")]
        return headers, values
    headers = ["Канал", "Дата", "Summary/текст", "Просмотры", "Репосты"]
    body = post.get("summary") or post.get("text") or ""
    values = [channel, str(post.get("posted_at") or ""), body,
              str(post.get("views") if post.get("views") is not None else ""),
              str(post.get("forwards") if post.get("forwards") is not None else "")]
    return headers, values


async def generate_report_rows(db_path: str, bot_id: int, since: str | None = None, channel_row_id: int | None = None):
    """Returns (headers, rows) — headers is the union-safe superset from the
    FIRST post's schema shape (channels with different schemas mixed in one
    report fall back gracefully: extra fields from later posts are dropped
    into the generic columns rather than crashing). Called by both the
    on-demand "📊 Отчёт" handler and the scheduled digest in
    runtime/userbot_worker.py, per the design's "no duplication" requirement."""
    posts = await get_posts_for_report(db_path, bot_id, since=since, channel_row_id=channel_row_id)
    if not posts:
        return [], []
    headers, first_row = _post_row(posts[0])
    rows = [first_row]
    for post in posts[1:]:
        h, row = _post_row(post)
        if h == headers:
            rows.append(row)
        else:
            # Schema mismatch across channels in the same report — degrade to
            # the generic 5-column shape rather than dropping the post.
            generic_headers = ["Канал", "Дата", "Summary/текст", "Просмотры", "Репосты"]
            body = post.get("summary") or post.get("text") or ""
            if not body and post.get("extracted"):
                # qa-stability review finding: json.loads on a stored `extracted`
                # value must degrade gracefully, not crash report generation —
                # same contract _post_row() already applies to this same column.
                try:
                    body = json.dumps(json.loads(post["extracted"]), ensure_ascii=False)
                except (ValueError, TypeError):
                    body = ""
            row = [post.get("channel_title") or post.get("channel_username") or "", str(post.get("posted_at") or ""),
                   body, str(post.get("views") or ""), str(post.get("forwards") or "")]
            if headers != generic_headers:
                # First post had a schema — keep original headers, pad this row.
                row = [row[0], row[1], row[2], *([""] * (len(headers) - 5)), row[3], row[4]]
            rows.append(row)
    return headers, rows


def _period_since(period: str) -> str | None:
    if period == "today":
        return datetime.now().strftime("%Y-%m-%d 00:00:00")
    if period == "week":
        return (datetime.now() - timedelta(days=7)).strftime("%Y-%m-%d %H:%M:%S")
    return None  # "all"


def build_csv_bytes(headers: list[str], rows: list[list[str]]) -> bytes:
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(headers)
    writer.writerows(rows)
    return buf.getvalue().encode("utf-8-sig")  # BOM so Excel opens Cyrillic correctly


def build_docx_bytes(headers: list[str], rows: list[list[str]], title: str = "Отчёт") -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def sync_report_to_sheets(bot_id: int, headers: list[str], rows: list[list[str]]) -> int:
    """Appends every row to the bot's connected Google Sheet via the shared
    sheets feature (features/sheets.write_row) — one worksheet 'Отчёт', header
    row written once if the worksheet doesn't exist yet (write_row creates it
    with add_worksheet, which starts empty, so we always write the header
    first when appending nothing existed). Raises ValueError if no sheet is
    connected for this bot_id — same contract as write_row itself, so callers
    can show a clear "подключите Google Sheets" message instead of a generic
    error. Returns the number of data rows written."""
    from features.sheets import write_row

    await write_row(bot_id, "Отчёт", headers)
    for row in rows:
        await write_row(bot_id, "Отчёт", row)
    return len(rows)


# ── handlers: entry ──────────────────────────────────────────────────────────
#
# SECURITY (docs audit 2026-08-19, project_multitenancy_audit_gaps memory):
# every handler in this router used to be reachable by ANY user in a private
# chat with the bot — no owner/admin check at all — letting an arbitrary
# stranger read the bot owner's monitoring feed, or worse, hijack the whole
# feature by authorizing the userbot session with their OWN phone number
# (the session is stored per bot_id, not per requesting user). Every handler
# below now requires _is_bot_admin() first, same posture as
# features/sellable_items.py's management handlers.

@router.message(F.text == "/monitor", F.chat.type == "private")
async def cmd_monitor(message: Message, state: FSMContext, bot_id: int, config: ChannelMonitorConfig):
    if not _is_bot_admin(message.from_user.id, config):
        return
    await state.clear()
    await message.answer(WELCOME_TEXT, parse_mode="HTML", reply_markup=kb_start_menu())


@router.callback_query(F.data == "cm_menu")
async def cb_menu(cb: CallbackQuery, state: FSMContext, bot_id: int, config: ChannelMonitorConfig):
    if not _is_bot_admin(cb.from_user.id, config):
        await cb.answer("⛔ Доступно только владельцу бота.", show_alert=True)
        return
    await cb.answer()
    await state.clear()
    await cb.message.edit_text(WELCOME_TEXT, parse_mode="HTML", reply_markup=kb_start_menu())


@router.callback_query(F.data == "cm_cancel")
async def cb_cancel(cb: CallbackQuery, state: FSMContext, bot_id: int, config: ChannelMonitorConfig):
    if not _is_bot_admin(cb.from_user.id, config):
        await cb.answer("⛔ Доступно только владельцу бота.", show_alert=True)
        return
    await cb.answer("Отменено")
    await state.clear()
    await cb.message.edit_text(WELCOME_TEXT, parse_mode="HTML", reply_markup=kb_start_menu())


# ── handlers: auth FSM ───────────────────────────────────────────────────────

@router.callback_query(F.data == "cm_connect")
async def cb_connect_start(cb: CallbackQuery, state: FSMContext, bot_id: int, config: ChannelMonitorConfig):
    """§6 of docs/USERBOT_CHANNEL_MONITOR_DESIGN.md: the risk screen must be
    shown BEFORE the phone number is ever requested — this is the entry
    point, not waiting_phone."""
    if not _is_bot_admin(cb.from_user.id, config):
        await cb.answer("⛔ Доступно только владельцу бота.", show_alert=True)
        return
    await cb.answer()
    await state.set_state(UserbotAuthStates.waiting_risk_ack)
    await cb.message.edit_text(RISK_TEXT, parse_mode="HTML", reply_markup=kb_risk_ack())


@router.callback_query(UserbotAuthStates.waiting_risk_ack, F.data == "cm_risk_ack")
async def cb_risk_ack(cb: CallbackQuery, state: FSMContext, bot_id: int, config: ChannelMonitorConfig):
    if not _is_bot_admin(cb.from_user.id, config):
        await cb.answer("⛔ Доступно только владельцу бота.", show_alert=True)
        return
    await cb.answer()
    await state.set_state(UserbotAuthStates.waiting_phone)
    await cb.message.edit_text(
        "Введите номер телефона аккаунта, который будет вести мониторинг, "
        "в международном формате (например +79991234567):",
        reply_markup=kb_cancel(),
    )


@router.message(UserbotAuthStates.waiting_phone, F.text, ~F.text.startswith("/"))
async def on_phone_entered(message: Message, state: FSMContext, bot_id: int, config: ChannelMonitorConfig):
    if not _is_bot_admin(message.from_user.id, config):
        await state.clear()
        return
    phone = message.text.strip()

    import time

    now = time.monotonic()
    cooldown_key = (bot_id, message.from_user.id)
    last = _last_phone_request_at.get(cooldown_key)
    if last is not None and (now - last) < PHONE_REQUEST_COOLDOWN_SECONDS:
        remaining = int(PHONE_REQUEST_COOLDOWN_SECONDS - (now - last))
        await message.answer(
            f"⏳ Подождите ещё {remaining} сек. перед следующей попыткой ввода номера.",
            reply_markup=kb_cancel(),
        )
        return
    _last_phone_request_at[cooldown_key] = now

    from telethon.errors import FloodWaitError, PhoneNumberInvalidError

    try:
        client = _make_client()
    except RuntimeError as e:
        logger.error(f"on_phone_entered: {e}")
        await state.clear()
        await message.answer(f"⚠️ {_esc(str(e), 300)}")
        return

    try:
        await client.connect()
        sent = await client.send_code_request(phone)
    except PhoneNumberInvalidError:
        await message.answer(
            "❌ Некорректный номер телефона. Введите номер в международном формате "
            "(например +79991234567):",
            reply_markup=kb_cancel(),
        )
        return
    except FloodWaitError as e:
        # Never retry sooner than Telegram says — see design doc §2.
        await state.clear()
        await message.answer(
            f"⏳ Telegram просит подождать {e.seconds} секунд перед следующей попыткой. "
            "Попробуйте снова позже.",
        )
        return
    except Exception as e:
        logger.error(f"on_phone_entered: send_code_request failed: {type(e).__name__}: {e}")
        await state.clear()
        await message.answer("⚠️ Не удалось начать авторизацию. Попробуйте позже.")
        return

    session_id = await create_userbot_session(config.db_path, bot_id, phone)
    await state.update_data(
        phone=phone,
        phone_code_hash=sent.phone_code_hash,
        session_id=session_id,
        code_attempts=0,
        session_string=client.session.save(),
    )
    await client.disconnect()
    await state.set_state(UserbotAuthStates.waiting_code)
    await message.answer(
        "📩 Код отправлен в ваш аккаунт Telegram. Введите его сюда:",
        reply_markup=kb_cancel(),
    )


async def _resume_client(state_data: dict):
    """Rebuilds a TelegramClient from the in-progress StringSession stashed in
    FSM state between turns — Telethon clients aren't picklable/storable
    directly in aiogram's FSM storage, so each turn reconnects using the
    session string captured after send_code_request/sign_in. This is only
    ever the in-progress (not yet 'active') session for one auth attempt."""
    client = _make_client(state_data.get("session_string", ""))
    await client.connect()
    return client


@router.message(UserbotAuthStates.waiting_code, F.text, ~F.text.startswith("/"))
async def on_code_entered(message: Message, state: FSMContext, config: ChannelMonitorConfig):
    if not _is_bot_admin(message.from_user.id, config):
        await state.clear()
        return
    from telethon.errors import (
        FloodWaitError, PhoneCodeExpiredError, PhoneCodeInvalidError, SessionPasswordNeededError,
    )

    code = message.text.strip()
    data = await state.get_data()
    phone = data.get("phone")
    phone_code_hash = data.get("phone_code_hash")
    session_id = data.get("session_id")
    attempts = data.get("code_attempts", 0)

    client = await _resume_client(data)
    try:
        await client.sign_in(phone, code, phone_code_hash=phone_code_hash)
    except PhoneCodeInvalidError:
        attempts += 1
        await client.disconnect()
        if attempts >= MAX_CODE_ATTEMPTS:
            await state.clear()
            if session_id:
                await set_userbot_session_status(config.db_path, session_id, "auth_failed")
            await message.answer("❌ Слишком много неверных попыток. Начните авторизацию заново.")
            return
        await state.update_data(code_attempts=attempts)
        await message.answer(
            f"❌ Неверный код ({attempts}/{MAX_CODE_ATTEMPTS}). Введите код ещё раз:",
            reply_markup=kb_cancel(),
        )
        return
    except PhoneCodeExpiredError:
        await client.disconnect()
        await state.set_state(UserbotAuthStates.waiting_phone)
        if session_id:
            await set_userbot_session_status(config.db_path, session_id, "auth_failed")
        await message.answer(
            "⌛ Код устарел. Введите номер телефона заново:",
            reply_markup=kb_cancel(),
        )
        return
    except SessionPasswordNeededError:
        await state.update_data(session_string=client.session.save())
        await client.disconnect()
        await state.set_state(UserbotAuthStates.waiting_2fa_password)
        await state.update_data(password_attempts=0)
        await message.answer(
            "🔐 На аккаунте включена облачная 2FA-защита. Введите пароль:",
            reply_markup=kb_cancel(),
        )
        return
    except FloodWaitError as e:
        await client.disconnect()
        await state.clear()
        if session_id:
            await set_userbot_session_status(config.db_path, session_id, "auth_failed")
        await message.answer(
            f"⏳ Telegram просит подождать {e.seconds} секунд перед следующей попыткой. "
            "Попробуйте снова позже.",
        )
        return
    except Exception as e:
        logger.error(f"on_code_entered: sign_in failed: {type(e).__name__}: {e}")
        await client.disconnect()
        await state.clear()
        if session_id:
            await set_userbot_session_status(config.db_path, session_id, "auth_failed")
        await message.answer("⚠️ Не удалось войти. Попробуйте позже.")
        return

    await _finish_auth_success(message, state, client, session_id, config)


@router.message(UserbotAuthStates.waiting_2fa_password, F.text, ~F.text.startswith("/"))
async def on_2fa_password_entered(message: Message, state: FSMContext, config: ChannelMonitorConfig):
    if not _is_bot_admin(message.from_user.id, config):
        await state.clear()
        return
    from telethon.errors import FloodWaitError, PasswordHashInvalidError

    password = message.text.strip()
    data = await state.get_data()
    session_id = data.get("session_id")
    attempts = data.get("password_attempts", 0)

    client = await _resume_client(data)
    try:
        await client.sign_in(password=password)
    except PasswordHashInvalidError:
        attempts += 1
        await client.disconnect()
        if attempts >= MAX_2FA_ATTEMPTS:
            await state.clear()
            if session_id:
                await set_userbot_session_status(config.db_path, session_id, "auth_failed")
            await message.answer("❌ Слишком много неверных попыток. Начните авторизацию заново.")
            return
        await state.update_data(password_attempts=attempts)
        await message.answer(
            f"❌ Неверный пароль ({attempts}/{MAX_2FA_ATTEMPTS}). Введите пароль ещё раз:",
            reply_markup=kb_cancel(),
        )
        return
    except FloodWaitError as e:
        await client.disconnect()
        await state.clear()
        if session_id:
            await set_userbot_session_status(config.db_path, session_id, "auth_failed")
        await message.answer(
            f"⏳ Telegram просит подождать {e.seconds} секунд перед следующей попыткой. "
            "Попробуйте снова позже.",
        )
        return
    except Exception as e:
        logger.error(f"on_2fa_password_entered: sign_in failed: {type(e).__name__}: {e}")
        await client.disconnect()
        await state.clear()
        if session_id:
            await set_userbot_session_status(config.db_path, session_id, "auth_failed")
        await message.answer("⚠️ Не удалось войти. Попробуйте позже.")
        return

    await _finish_auth_success(message, state, client, session_id, config)


async def _finish_auth_success(
    message: Message, state: FSMContext, client, session_id: int | None, config: ChannelMonitorConfig
) -> None:
    """§2 step 4 of docs/USERBOT_CHANNEL_MONITOR_DESIGN.md: session.save() →
    encrypted → DB, status='active'. Never logs or echoes the session string
    itself."""
    session_string = client.session.save()
    await client.disconnect()
    await state.clear()
    if session_id:
        await set_userbot_session_active(config.db_path, session_id, session_string)
    await message.answer(
        "✅ Мониторинг подключён! Теперь можно добавить каналы для отслеживания.\n\n"
        + RISK_TEXT,
        parse_mode="HTML",
        reply_markup=kb_start_menu(),
    )


# ── handlers: channel list / toggle / feed ───────────────────────────────────

@router.callback_query(F.data == "cm_channels")
async def cb_channels(cb: CallbackQuery, state: FSMContext, bot_id: int, config: ChannelMonitorConfig):
    if not _is_bot_admin(cb.from_user.id, config):
        await cb.answer("⛔ Доступно только владельцу бота.", show_alert=True)
        return
    await cb.answer()
    await state.clear()
    text, channels = await _channels_panel_text(config.db_path, bot_id)
    await cb.message.edit_text(text, parse_mode="HTML", reply_markup=kb_channels(channels))


@router.callback_query(F.data.startswith("cm_toggle:"))
async def cb_toggle_channel(cb: CallbackQuery, bot_id: int, config: ChannelMonitorConfig):
    if not _is_bot_admin(cb.from_user.id, config):
        await cb.answer("⛔ Доступно только владельцу бота.", show_alert=True)
        return
    channel_id = int(cb.data.split(":", 1)[1])
    channels = await get_monitored_channels_by_bot(config.db_path, bot_id)
    target = next((c for c in channels if c["id"] == channel_id), None)
    if not target:
        await cb.answer("Канал не найден", show_alert=True)
        return
    new_active = not target.get("active")
    await set_monitored_channel_active(config.db_path, channel_id, new_active)
    await cb.answer("Включено" if new_active else "Выключено")
    text, channels = await _channels_panel_text(config.db_path, bot_id)
    await cb.message.edit_text(text, parse_mode="HTML", reply_markup=kb_channels(channels))


@router.callback_query(F.data == "cm_add_channel")
async def cb_add_channel_start(cb: CallbackQuery, state: FSMContext, bot_id: int, config: ChannelMonitorConfig):
    if not _is_bot_admin(cb.from_user.id, config):
        await cb.answer("⛔ Доступно только владельцу бота.", show_alert=True)
        return
    await cb.answer()
    await state.set_state(_AddChannelState.waiting_username)
    await cb.message.edit_text(
        "Введите @username канала, который нужно отслеживать:",
        reply_markup=kb_cancel(),
    )


@router.message(_AddChannelState.waiting_username, F.text, ~F.text.startswith("/"))
async def on_channel_username_entered(message: Message, state: FSMContext, bot_id: int, config: ChannelMonitorConfig):
    """Resolves the given @username via THIS bot's active userbot client and
    joins it if public — per docs/USERBOT_CHANNEL_MONITOR_DESIGN.md §4
    (Telegram's Client-API constraint that a userbot only ever sees channels
    its own account has joined applies identically regardless of which bot
    hosts this feature — see docs/USERBOT_FEATURE_DESIGN.md §5). A private/
    inaccessible channel gets an explicit error, never a silent no-op."""
    if not _is_bot_admin(message.from_user.id, config):
        await state.clear()
        return
    username = message.text.strip().lstrip("@")
    await state.clear()

    sessions = await get_userbot_sessions_by_bot(config.db_path, bot_id)
    active_session = next((s for s in sessions if s.get("status") == "active" and s.get("session_string")), None)
    if not active_session:
        await message.answer(
            "⚠️ Сначала подключите мониторинг («➕ Подключить мониторинг»), затем добавляйте каналы.",
            reply_markup=kb_back_to_menu(),
        )
        return

    # Owner decision #4: hard limit of 50 channels per bot — checked BEFORE
    # attempting the join, so a bot at the cap never even touches Telegram.
    current_count = await count_monitored_channels(config.db_path, bot_id)
    if current_count >= MAX_CHANNELS_PER_BOT:
        await message.answer(
            f"🚫 Достигнут лимит в {MAX_CHANNELS_PER_BOT} каналов на бота. Отключите "
            "неиспользуемый канал в «📋 Мои каналы», чтобы добавить новый.",
            reply_markup=kb_back_to_menu(),
        )
        return

    from telethon.errors import ChannelPrivateError, UsernameInvalidError, UsernameNotOccupiedError
    from telethon.tl.functions.channels import JoinChannelRequest

    client = _make_client(active_session["session_string"])
    try:
        await client.connect()
        entity = await client.get_entity(username)
        await client(JoinChannelRequest(entity))
    except (ChannelPrivateError,):
        await client.disconnect()
        await message.answer(
            "🔒 Этот канал закрытый — мониторинг возможен только если ваш аккаунт уже "
            "состоит в нём или получит приглашение отдельно. Работает автоматически "
            "только для публичных каналов (по @username). Для закрытого канала "
            "попросите администратора добавить ваш аккаунт вручную, или получите и "
            "самостоятельно перейдите по инвайт-ссылке (бот не вступает по ссылке "
            "автоматически). Подробнее — кнопка «❓ Как добавить закрытый канал» в меню.",
            reply_markup=kb_back_to_menu(),
        )
        return
    except (UsernameInvalidError, UsernameNotOccupiedError):
        await client.disconnect()
        await message.answer(
            "❌ Канал с таким @username не найден. Проверьте имя и попробуйте ещё раз.",
            reply_markup=kb_back_to_menu(),
        )
        return
    except Exception as e:
        logger.error(f"on_channel_username_entered: resolve/join failed: {type(e).__name__}: {e}")
        await client.disconnect()
        await message.answer(
            "⚠️ Не удалось добавить канал (закрытый, недоступен, или временная ошибка "
            "Telegram). Попробуйте позже.",
            reply_markup=kb_back_to_menu(),
        )
        return

    channel_id = getattr(entity, "id", None)
    channel_title = getattr(entity, "title", None) or username
    await client.disconnect()
    channel_row_id = await add_monitored_channel(config.db_path, bot_id, username, channel_id, channel_title)
    await state.update_data(pending_channel_row_id=channel_row_id, pending_channel_title=channel_title)
    await state.set_state(_AddChannelState.waiting_schema_choice)
    await message.answer(
        f"✅ Канал <b>{_esc(channel_title)}</b> добавлен и будет отслеживаться.\n\n"
        "Теперь выберите, что извлекать из его постов — готовый пресет полей, "
        "свою схему, либо оставить только общий summary без схемы:",
        parse_mode="HTML",
        reply_markup=kb_schema_picker(),
    )


@router.callback_query(_AddChannelState.waiting_schema_choice, F.data.startswith("cm_schema:"))
async def cb_schema_choice(cb: CallbackQuery, state: FSMContext, config: ChannelMonitorConfig):
    if not _is_bot_admin(cb.from_user.id, config):
        await cb.answer("⛔ Доступно только владельцу бота.", show_alert=True)
        return
    key = cb.data.split(":", 1)[1]
    data = await state.get_data()
    channel_row_id = data.get("pending_channel_row_id")
    if not channel_row_id:
        await cb.answer()
        await state.clear()
        await cb.message.edit_text(WELCOME_TEXT, parse_mode="HTML", reply_markup=kb_start_menu())
        return

    if key == "custom":
        await cb.answer()
        await state.set_state(_AddChannelState.waiting_custom_fields)
        await cb.message.edit_text(
            "Перечислите названия полей через запятую (например: цена, срок, площадь):",
            reply_markup=kb_cancel(),
        )
        return

    await cb.answer()
    if key == "none":
        await set_monitored_channel_schema(config.db_path, channel_row_id, None)
    else:
        preset = _PRESETS_BY_KEY.get(key)
        fields = preset["fields"] if preset else []
        await set_monitored_channel_schema(config.db_path, channel_row_id, json.dumps(fields, ensure_ascii=False))
    await state.clear()
    await cb.message.edit_text(
        "✅ Схема сохранена. Канал будет отслеживаться с выбранными настройками извлечения.",
        reply_markup=kb_back_to_menu(),
    )


@router.message(_AddChannelState.waiting_custom_fields, F.text, ~F.text.startswith("/"))
async def on_custom_fields_entered(message: Message, state: FSMContext, config: ChannelMonitorConfig):
    if not _is_bot_admin(message.from_user.id, config):
        await state.clear()
        return
    data = await state.get_data()
    channel_row_id = data.get("pending_channel_row_id")
    fields = [f.strip() for f in message.text.split(",") if f.strip()]
    if not fields:
        await message.answer(
            "⚠️ Не удалось разобрать список полей. Введите хотя бы одно название через запятую:",
            reply_markup=kb_cancel(),
        )
        return
    await state.clear()
    if channel_row_id:
        await set_monitored_channel_schema(config.db_path, channel_row_id, json.dumps(fields, ensure_ascii=False))
    await message.answer(
        f"✅ Свободная схема сохранена: {', '.join(fields)}.",
        reply_markup=kb_back_to_menu(),
    )


@router.callback_query(F.data == "cm_closed_guide")
async def cb_closed_guide(cb: CallbackQuery, config: ChannelMonitorConfig):
    if not _is_bot_admin(cb.from_user.id, config):
        await cb.answer("⛔ Доступно только владельцу бота.", show_alert=True)
        return
    await cb.answer()
    await cb.message.edit_text(CLOSED_CHANNEL_GUIDE_TEXT, parse_mode="HTML", reply_markup=kb_back_to_menu())


# ── handlers: report (on-demand "📊 Отчёт") ──────────────────────────────────

@router.callback_query(F.data == "cm_report")
async def cb_report_start(cb: CallbackQuery, state: FSMContext, config: ChannelMonitorConfig):
    if not _is_bot_admin(cb.from_user.id, config):
        await cb.answer("⛔ Доступно только владельцу бота.", show_alert=True)
        return
    await cb.answer()
    await state.clear()
    await cb.message.edit_text("Выберите период отчёта:", reply_markup=kb_report_periods())


@router.callback_query(F.data.startswith("cm_report_period:"))
async def cb_report_period(cb: CallbackQuery, state: FSMContext, bot_id: int, config: ChannelMonitorConfig):
    if not _is_bot_admin(cb.from_user.id, config):
        await cb.answer("⛔ Доступно только владельцу бота.", show_alert=True)
        return
    period = cb.data.split(":", 1)[1]
    await state.update_data(report_period=period)
    channels = await get_monitored_channels_by_bot(config.db_path, bot_id)
    await cb.answer()
    await state.set_state(_ReportState.waiting_channel_choice)
    await cb.message.edit_text("Выберите канал (или все сразу):", reply_markup=kb_report_channels(channels))


@router.callback_query(_ReportState.waiting_channel_choice, F.data.startswith("cm_report_channel:"))
async def cb_report_channel(cb: CallbackQuery, state: FSMContext, config: ChannelMonitorConfig):
    if not _is_bot_admin(cb.from_user.id, config):
        await cb.answer("⛔ Доступно только владельцу бота.", show_alert=True)
        return
    choice = cb.data.split(":", 1)[1]
    await state.update_data(report_channel=None if choice == "all" else int(choice))
    await cb.answer()
    await cb.message.edit_text("В каком формате прислать отчёт?", reply_markup=kb_report_formats())


@router.callback_query(F.data.startswith("cm_report_fmt:"))
async def cb_report_format(cb: CallbackQuery, state: FSMContext, bot_id: int, config: ChannelMonitorConfig):
    if not _is_bot_admin(cb.from_user.id, config):
        await cb.answer("⛔ Доступно только владельцу бота.", show_alert=True)
        return
    fmt = cb.data.split(":", 1)[1]
    data = await state.get_data()
    period = data.get("report_period", "all")
    channel_row_id = data.get("report_channel")
    await state.clear()
    await cb.answer()

    since = _period_since(period)
    headers, rows = await generate_report_rows(config.db_path, bot_id, since=since, channel_row_id=channel_row_id)
    if not rows:
        await cb.message.edit_text("За выбранный период постов нет.", reply_markup=kb_back_to_menu())
        return

    if fmt == "csv":
        content = build_csv_bytes(headers, rows)
        await cb.message.answer_document(
            BufferedInputFile(content, filename="report.csv"),
            caption=f"📊 Отчёт ({REPORT_PERIODS.get(period, period)}), {len(rows)} строк.",
        )
    elif fmt == "docx":
        content = build_docx_bytes(headers, rows)
        await cb.message.answer_document(
            BufferedInputFile(content, filename="report.docx"),
            caption=f"📊 Отчёт ({REPORT_PERIODS.get(period, period)}), {len(rows)} строк.",
        )
    elif fmt == "sheets":
        try:
            n = await sync_report_to_sheets(bot_id, headers, rows)
        except ValueError:
            await cb.message.edit_text(
                "⚠️ Google Sheets не подключены для этого бота. Подключите таблицу в "
                "настройках бота, затем попробуйте снова.",
                reply_markup=kb_back_to_menu(),
            )
            return
        except Exception as e:
            logger.error(f"cb_report_format: sheets sync failed: {type(e).__name__}: {e}")
            await cb.message.edit_text(
                "⚠️ Не удалось записать отчёт в Google Sheets. Попробуйте позже.",
                reply_markup=kb_back_to_menu(),
            )
            return
        await cb.message.answer(f"✅ Записано {n} строк в лист «Отчёт» вашей Google-таблицы.")
        return

    await cb.message.answer("Готово.", reply_markup=kb_back_to_menu())


# ── handlers: report schedule ("⏰ Настроить расписание") ────────────────────

@router.callback_query(F.data == "cm_schedule")
async def cb_schedule_start(cb: CallbackQuery, state: FSMContext, config: ChannelMonitorConfig):
    if not _is_bot_admin(cb.from_user.id, config):
        await cb.answer("⛔ Доступно только владельцу бота.", show_alert=True)
        return
    await cb.answer()
    await state.clear()
    await cb.message.edit_text("Как часто присылать отчёт?", reply_markup=kb_schedule_frequency())


@router.callback_query(F.data.startswith("cm_freq:"))
async def cb_schedule_frequency(cb: CallbackQuery, state: FSMContext, config: ChannelMonitorConfig):
    if not _is_bot_admin(cb.from_user.id, config):
        await cb.answer("⛔ Доступно только владельцу бота.", show_alert=True)
        return
    frequency = cb.data.split(":", 1)[1]
    await state.update_data(schedule_frequency=frequency)
    await cb.answer()
    await state.set_state(_ScheduleState.waiting_time)
    await cb.message.edit_text(
        "Введите время отправки в формате ЧЧ:ММ (например 09:00):",
        reply_markup=kb_cancel(),
    )


_TIME_RE = re.compile(r"^([01]?\d|2[0-3]):([0-5]\d)$")


@router.message(_ScheduleState.waiting_time, F.text, ~F.text.startswith("/"))
async def on_schedule_time_entered(message: Message, state: FSMContext, bot_id: int, config: ChannelMonitorConfig):
    if not _is_bot_admin(message.from_user.id, config):
        await state.clear()
        return
    m = _TIME_RE.match(message.text.strip())
    if not m:
        await message.answer(
            "⚠️ Неверный формат. Введите время как ЧЧ:ММ (например 09:00):",
            reply_markup=kb_cancel(),
        )
        return
    hour, minute = int(m.group(1)), int(m.group(2))
    data = await state.get_data()
    frequency = data.get("schedule_frequency", "daily")
    if frequency == "weekly":
        await state.update_data(schedule_hour=hour, schedule_minute=minute)
        await state.set_state(_ScheduleState.waiting_weekday)
        await message.answer(
            "Введите день недели числом (1 — понедельник, ..., 7 — воскресенье):",
            reply_markup=kb_cancel(),
        )
        return

    await set_report_schedule(config.db_path, bot_id, "daily", hour, minute)
    await state.clear()
    await message.answer(
        f"✅ Расписание сохранено: каждый день в {hour:02d}:{minute:02d}.",
        reply_markup=kb_back_to_menu(),
    )


@router.message(_ScheduleState.waiting_weekday, F.text, ~F.text.startswith("/"))
async def on_schedule_weekday_entered(message: Message, state: FSMContext, bot_id: int, config: ChannelMonitorConfig):
    if not _is_bot_admin(message.from_user.id, config):
        await state.clear()
        return
    raw = message.text.strip()
    if not raw.isdigit() or not (1 <= int(raw) <= 7):
        await message.answer(
            "⚠️ Введите число от 1 (понедельник) до 7 (воскресенье):",
            reply_markup=kb_cancel(),
        )
        return
    weekday = int(raw) - 1  # 0=Monday, matches date.weekday()
    data = await state.get_data()
    hour, minute = data.get("schedule_hour"), data.get("schedule_minute")
    if hour is None or minute is None:
        # qa-stability review finding: FSM state can be lost mid-wizard (e.g. a
        # process restart between turns, same known class of bug as
        # backlog_create_flow_pending_race) — report_schedules.hour/minute are
        # NOT NULL, so inserting None would crash with an uncaught
        # aiosqlite.IntegrityError. Fail soft: ask the owner to restart the
        # schedule wizard instead of letting the handler crash.
        await state.clear()
        await message.answer(
            "⚠️ Не удалось определить время — начните настройку расписания заново "
            "(«⏰ Настроить расписание»).",
            reply_markup=kb_back_to_menu(),
        )
        return
    await set_report_schedule(config.db_path, bot_id, "weekly", hour, minute, weekday=weekday)
    await state.clear()
    await message.answer(
        f"✅ Расписание сохранено: каждую {WEEKDAY_NAMES[weekday]} в {hour:02d}:{minute:02d}.",
        reply_markup=kb_back_to_menu(),
    )


@router.callback_query(F.data == "cm_feed")
async def cb_feed(cb: CallbackQuery, state: FSMContext, bot_id: int, config: ChannelMonitorConfig):
    if not _is_bot_admin(cb.from_user.id, config):
        await cb.answer("⛔ Доступно только владельцу бота.", show_alert=True)
        return
    await cb.answer()
    await state.clear()
    text = await _feed_text(config.db_path, bot_id)
    await cb.message.edit_text(text, parse_mode="HTML", reply_markup=kb_back_to_menu())


# ── handlers: revoke ─────────────────────────────────────────────────────────

@router.message(F.text == "/disconnect_monitor", F.chat.type == "private")
async def cmd_disconnect(message: Message, state: FSMContext, bot_id: int, config: ChannelMonitorConfig):
    if not _is_bot_admin(message.from_user.id, config):
        return
    await state.clear()
    sessions = await get_userbot_sessions_by_bot(config.db_path, bot_id)
    active = [s for s in sessions if s.get("status") == "active"]
    if not active:
        await message.answer("У вас нет активного подключения мониторинга.")
        return
    for s in active:
        await revoke_userbot_session(config.db_path, s["id"])
    await message.answer(
        "🔴 Мониторинг отключён. Сессия удалена — для повторного подключения потребуется "
        "полная авторизация заново.",
    )
